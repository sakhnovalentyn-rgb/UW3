from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import gc
import os
import re
import io
import json
import glob
import unicodedata
import pandas as pd
from dateutil.relativedelta import relativedelta
import traceback
import colorama
from colorama import init, Fore, Style
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from streamlit import metric
init(autoreset=True)
from datetime import datetime

def clean_and_deduplicate(text, seen_set):
    """
    Чистить текст від сміття та перевіряє, чи не бачили ми його раніше.
    Повертає очищений текст або None, якщо це дублікат.
    """
    if not text or text == "Дані відсутні.":
        return None
    
    # Базова чистка
    text = re.sub(r"^.*?–\s*", "", text).strip()
    
    # Створюємо "відбиток" тексту (перші 60 символів без пробілів та регістру)
    fingerprint = "".join(text[:60].lower().split())
    
    if fingerprint in seen_set or not fingerprint:
        return None
        
    seen_set.add(fingerprint)
    return text

def insert_styled_text(paragraph, text):
    """Додає текст у Word. Якщо даних немає — фарбує в червоний."""
    run = paragraph.add_run(str(text))
    if text == "ДАНІ НЕ ЗНАЙДЕНО" or not text or str(text).strip() == "":
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
    return run

def consume_part(pool, keywords, stop_markers):
    """Шукає шматок тексту в пулі, повертає його і видаляє з пулу."""
    if not pool or pool == "ДАНІ НЕ ЗНАЙДЕНО":
        return "ДАНІ НЕ ЗНАЙДЕНО", pool
    
    # Екрануємо символи, щоб "." чи "(" не зламали пошук
    escaped_keys = [re.escape(k) for k in keywords]
    escaped_stops = [re.escape(s) for s in stop_markers]
    
    pattern = rf"({'|'.join(escaped_keys)})[\s:–-]*(.*?)(?={'|'.join(escaped_stops)}|$)"
    
    match = re.search(pattern, pool, re.IGNORECASE | re.DOTALL)
    
    if match:
        content = match.group(2).strip()
        full_match = match.group(0)
        
        # Чистимо від розділювачів таблиць та зайвих пробілів
        content = content.replace('|', ' ')
        content = re.sub(r'\s+', ' ', content).strip()
        
        # Тут можна додати вашу функцію clean_bank_jargon(content), якщо вона є
        new_pool = pool.replace(full_match, "").strip()
        return content if content else "ДАНІ НЕ ЗНАЙДЕНО", new_pool
    
    return "ДАНІ НЕ ЗНАЙДЕНО", pool

def clean_bank_jargon(text):
    """Базова чистка тексту від технічного сміття анкети."""
    if not text: return text
    # Видаляємо повторювані заголовки, які часто залітають в текст
    text = re.sub(r'^(Так/Ні|Дані відсутні|Заповнювати лише для|Обов’язково зазначається).*?$', '', text, flags=re.MULTILINE)
    return text.strip()

def set_cell_background(cell, fill_color):
    """Встановлює колір фону комірки (fill_color у форматі HEX, наприклад 'FFFF00')"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def clean_key(key):
    """Normalize and clean a string key."""
    key = unicodedata.normalize("NFC", key).replace('\xa0', ' ').replace('’', "'").replace('\u2019', "'").replace('\u2018', "'")
    return re.sub(r'\s+', ' ', key).strip()

alias_map = {
    # Основні фінанси
    "ebitda": ["EBITDA", "ЕБІТДА"],
    "revenue": ["Виручка від реалізації", "Чистий дохід", "Дохід (виручка)"],
    "costs": ["Собівартість"],
    "current_assets": ["Всього обігові", "Оборотні активи"],
    "short_term_liabilities": ["Короткострокові зобов'язання", "Поточні зобов'язання"],
    "total_balance": ["Всього баланс", "Баланс"],
    "equity": ["Власний капітал"],
    "fixed_assets": ["ОЗ", "Основні засоби"],
    "profit_loss": ["Прибуток (збиток)", "Чистий прибуток"],
    
    # Кредит та платежі
    "total_requested_credit": ["Запитуваний кредит:", "Сума кредиту", "Загальна сума запиту"],
    "total_annual_payment": ["Річний платіж", "Total annual payment"],
    "total_annual_payments_other_banks": ["Платежі в інших банках", "Інші банки платежі"],
    "total_annual_payments_history": ["Історія платежів", "Погашення за історією"],
    
    # Специфічні банківські показники
    "total_inflow": ["Всього надходження", "Обороти по рахунку", "total_inflow"],
    "turnover_12m": ["Оборот за 12 місяців", "Виручка за рік"],
    "IF_otherbanks_total": ["IF інші банки", "Інвестиційні кредити інші"],
    "WCF_otherbanks_total": ["WCF інші банки", "Обігові кошти інші"],
    "IF_RBUA": ["IF RBUA", "Кредити інвестиційні Райффайзен"],
    "WCF_RBUA": ["WCF RBUA", "Кредити обігові Райффайзен"],
    
    # Інше
    "kved_code": ["КВЕД"],
    "borrower_status": ["Інформація щодо реального стану"],
    "scoring_result": ["URG", "Клас позичальника"]
}
current_dir = os.path.dirname(os.path.abspath(__file__))
rules_path = os.path.join(current_dir, 'rules.json')

try:
    with open(rules_path, 'r', encoding='utf-8') as f:
        rules = json.load(f)
    print(f"✅ Rules loaded successfully from {rules_path}")
except Exception as e:
    print(f"⚠️ Error loading rules.json: файл не знайдено за шляхом {rules_path}")
    text_rules = {"Плюси": [], "Мінуси": [], "Ризики": [], "+/-": []}
    exit(1)

from docx.oxml.shared import qn
from docx.oxml import OxmlElement

def set_table_borders(table):
    """Примусово додає чорні межі до всіх клітинок таблиці через XML."""
    tbl = table._tbl
    # Перевіряємо наявність властивостей таблиці (tblPr)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Шукаємо або створюємо елемент меж (tblBorders)
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    # Додаємо всі 6 типів ліній (зовнішні та внутрішні)
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        
        # Налаштування: одинарна лінія, товщина 4 (0.5 пт), чорний колір
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')

def normalize_financial_data(financial_data):
    """Normalize financial data, converting thousands to UAH."""
    print(f"📝 normalize_financial_data: input: {financial_data}")
    normalized_data = {}
    numeric_keys = ["ebitda", "revenue", "costs", "current_assets", "short_term_liabilities", "equity", "total_balance", "fixed_assets", "profit_loss", "total_requested_credit","total_annual_payment", "total_inflow"]
    for key, value in financial_data.items():
        if key in numeric_keys:
            try:
                cleaned_value = str(value).replace(" ", "").replace(",", ".").replace("грн", "").strip()
                normalized_value = float(cleaned_value)
                if normalized_value < 1000000:
                    normalized_value *= 1000
                    print(f"✅ Converted {key} from thousands to UAH: {value} -> {normalized_value}")
                normalized_data[key] = normalized_value
                print(f"✅ Normalized {key}: {value} -> {normalized_data[key]} UAH")
            except (ValueError, TypeError):
                print(f"⚠️ Failed to normalize {key}: value {value} is not numeric")
                normalized_data[key] = 0.0
        elif key in ["kved_code", "borrower_status", "credits", "credit_history_payments"]:
            normalized_data[key] = value
            print(f"✅ Preserved {key}: {value}")
        else:
            print(f"📝 Skipping key {key}, not in expected keys")
    print(f"📝 normalize_financial_data: output: {normalized_data}")
    return normalized_data

def get_text_comments(text, rules_list):
    found_comments = []
    if not text:
        return found_comments
    
    # Очищуємо вхідний текст від зайвих переносів рядків для легшого пошуку
    clean_text = " ".join(text.split())
    
    for rule in rules_list:
        pattern = None
        if 'pattern' in rule:
            pattern = rule['pattern']
        elif 'text_extraction' in rule and isinstance(rule['text_extraction'], dict):
            pattern = rule['text_extraction'].get('pattern')

        if pattern:
            # ГОЛОВНЕ: Прибираємо суворі прив'язки \s* та ^ $ з паттерна
            # Це дозволяє знайти "Майно підприємства" навіть якщо в реченні є інші слова
            clean_pattern = pattern.replace('\\s*', '').replace('^', '').replace('$', '').strip()
            
            try:
                if re.search(clean_pattern, clean_text, re.IGNORECASE):
                    # Якщо в правилі є 'comment', беремо його, якщо ні (для Ризиків) - беремо з thresholds
                    comment = rule.get('comment')
                    if not comment and 'thresholds' in rule:
                        comment = rule['thresholds'][0].get('comment')
                    
                    if comment:
                        found_comments.append(comment)
            except Exception as e:
                print(f"Regex error: {e}")
    
    return found_comments
def extract_scoring_result(text):
    if not text:
        return None
    # Шукаємо фразу "Результат скорингу:" та забираємо все до кінця рядка
    match = re.search(r"Результат скорингу:\s*([\d,.]+)", str(text), re.IGNORECASE)
    if match:
        return match.group(1)  # Повертаємо текст точно так, як він у файлі
    return None
    
def extract_financial_metrics_from_json(json_data, alias_map):
    financial_data = {}
    clean_pattern = r'[^\d,.-]'
    loan_amount, loan_term, loan_rate = 0, 0, 0

    def process_table(table_obj):
        nonlocal loan_amount, loan_term, loan_rate
        
        for row in table_obj.get("rows", []):
            cells = row.get("cells", [])
            if not cells: continue
            
            label = str(cells[0]).strip()
            
            for nested in row.get("nested_tables", []):
                process_table(nested)

            if "Сумма" in label and len(cells) >= 4:
                pass 
            
            if any("UAH" in str(c) for c in cells) and len(cells) >= 4:
                try:
                    loan_amount = float(str(cells[0]).replace(' ', '').replace(',', '.'))
                    loan_term = float(str(cells[2]).replace(' ', '').replace(',', '.'))
                    loan_rate = float(str(cells[3]).replace(' ', '').replace(',', '.'))
                    if loan_rate > 1: loan_rate /= 100
                    
                    financial_data["total_requested_credit"] = loan_amount
                    print(f"✅ ЗНАЙДЕНО В NESTED: Кредит={loan_amount}, Термін={loan_term}, Ставка={loan_rate}")
                except:
                    pass

            # СТАНДАРТНИЙ ПОШУК ЗА АЛІАСАМИ
            label_lower = label.lower()
            for target_key, aliases in alias_map.items():
                if any(alias.lower() in label_lower for alias in aliases):
                    if target_key not in financial_data and len(cells) > 1:
                        val_str = str(cells[1]).strip()
                        cleaned = re.sub(clean_pattern, '', val_str).replace(',', '.')
                        try:
                            financial_data[target_key] = float(cleaned)
                        except: continue

    # Запускаємо обробку для всіх основних таблиць
    for main_table in json_data.get("tables", []):
        process_table(main_table)

#    ФІНАЛЬНИЙ РОЗРАХУНОК ПЛАТЕЖУ
    if loan_amount > 0:
        # Обмеження ставки: не більше 30%
        # Спочатку перевіряємо, чи не передана ставка як ціле число (наприклад, 25 замість 0.25)
        current_rate = loan_rate if loan_rate < 1 else loan_rate / 100
        
        # Застосовуємо ліміт 30% (0.30)
        final_rate = min(current_rate, 0.30)
        
        term_calc = loan_term if loan_term > 0 else 12
        
        # Розрахунок з урахуванням обмеженої ставки
        annual_pay = (loan_amount / term_calc * 12) + (loan_amount * final_rate)
        
        financial_data["total_annual_payment"] = annual_pay
        
        if final_rate < current_rate:
            print(f"⚠️ СТАВКА ОБМЕЖЕНА: {current_rate*100}% -> {final_rate*100}%")
            
        print(f"🧮 РОЗРАХОВАНО ПЛАТЕЖ: {annual_pay:,.2f}")

    return financial_data

def process_table(table, processed_tables):
    credit_pattern = r"\b(\d{1,3}(?:\s\d{3})*,\d{0,2})\s*(?:грн\.?|UAH)?\b"
    processed_credits = set()

    table_id = table.get("table_id", f"Table_{len(processed_tables) + 1}")
    if table_id in processed_tables:
        print(f"📝 Table {table_id} already processed, skipping")
        return
    processed_tables.add(table_id)
    print(f"📝 Found table: {table_id}")
    
    rows = table.get("rows", [])
    print(f"📝 Number of rows in table {table_id}: {len(rows)}")
    
    if table_id == "Table_12.1.11":
        for row_idx, row in enumerate(rows, start=1):
            cells = row.get("cells", [])
            print(f"📝 Row {row_idx} in table {table_id}: {cells}")
            if len(cells) >= 2:
                label = clean_key(cells[0])
                value = cells[1].strip()
                alias = alias_map.get(label)
                if alias:
                    try:
                        cleaned_value = value.replace(" ", "").replace(",", ".").rstrip("\n")
                        metric[alias] = float(cleaned_value)
                        print(f"✅ Extracted {alias}: {metric[alias]}")
                    except ValueError:
                        print(f"⚠️ Cannot convert value for {label}: {value}")
    
    # Process credit tables (e.g., Table_1.11.11 or nested tables)
    if any(table_id.startswith(f"Table_{i}.") for i in range(1, 13)) or table_id in ["Table_1.11.11", "Table_1.12.11", "Table_20"]:
        for row_idx, row in enumerate(rows, start=1):
            cells = row.get("cells", [])
            print(f"📝 Row {row_idx} in table {table_id}: {cells}")
            if len(cells) >= 5 and row_idx == 2:  # Second row contains credit data
                credit_info = {"amount": 0.0, "type": "Невідомий", "rate": None, "term": None}
                try:
                    amount_str = cells[0].strip()
                    print(f"📝 Extracting amount: {amount_str}")
                    if re.match(credit_pattern, amount_str):
                        cleaned_amount = amount_str.replace(" ", "").replace(",", ".").replace("грн", "").strip()
                        credit_amount = float(cleaned_amount)*1000
                        credit_id = f"{table_id}_{credit_amount}"
                        if credit_id in processed_credits:
                            print(f"📝 Skipping duplicate amount: {credit_amount}")
                            continue
                        processed_credits.add(credit_id)
                        credit_info["amount"] = credit_amount
                        metric["total_requested_credit"] += credit_amount
                        print(f"✅ Added amount: {credit_amount}, total_requested_credit: {metric['total_requested_credit']}")
                    term_str = cells[2].strip()
                    credit_info["term"] = int(term_str) if term_str and term_str.isdigit() and int(term_str) > 0 else 12
                    print(f"✅ Term set: {credit_info['term']}")
                    rate_str = cells[3].strip()
                    try:
                        cleaned_rate = rate_str.replace(" ", "").replace(",", ".").replace("%", "").strip()
                        rate = float(cleaned_rate) / 100
                        credit_info["rate"] = rate if 0 <= rate <= 1 else 0.20
                        print(f"✅ Rate set: {credit_info['rate']}")
                    except ValueError:
                        credit_info["rate"] = 0.20
                        print(f"📝 Default rate used: 20%")
                    product = cells[4].strip()
                    if product:
                        product_lower = product.lower()
                        known_types = ["овердрафт", "вкл", "довіра", "розвиток", "інвест"]
                        credit_info["type"] = next(
                            (keyword.capitalize() for keyword in known_types if keyword in product_lower),
                            product.split()[0] if product.split() else "Невідомий"
                        )
                        print(f"✅ Product type set: {credit_info['type']}")
                    if credit_info["amount"] > 0:
                        metric["credits"].append(credit_info)
                        print(f"✅ Added credit to metrics['credits']: {credit_info}")
                except Exception as e:
                    print(f"⚠️ Error processing row {row_idx} in table {table_id}: {e}")
    
#        # Process KVED and other metrics
#        for row in rows:
#            row_text = " | ".join(row.get("cells", []))
#            row_text = unicodedata.normalize("NFC", row_text).replace('\xa0', ' ').replace('’', "'").replace('\u2019', "'").replace('\u2018', "'")
#            kved_matches = re.findall(kved_pattern, row_text, re.UNICODE)
#            if kved_matches:
#                metrics["kved_code"] = kved_matches[0].strip()
#                print(f"✅ Extracted KVED: {metrics['kved_code']}")
#            matches = re.findall(pattern, row_text, re.UNICODE | re.MULTILINE)
#            for label, value in matches:
#                label_clean = clean_key(label)
#                alias = alias_map.get(label_clean)
#                if alias:
#                    try:
#                        if alias == "kved_code":
#                           metrics[alias] = value.strip()
#                            print(f"✅ Extracted {alias}: {metrics[alias]}")
#                        else:
#                            cleaned_value = value.replace(" ", "").replace(",", ".").rstrip("\n")
#                            metrics[alias] = float(cleaned_value)
#                            print(f"✅ Extracted {alias}: {metrics[alias]}")
#                    except ValueError:
#                        print(f"⚠️ Cannot convert value for {label_clean}: {value}")
#        
#        # Process nested tables
#        for row in rows:
#            for nested_table in row.get("nested_tables", []):
#                process_table(nested_table, processed_tables)
#    
#    processed_tables = set()
#    for table in json_data.get("tables", []):
#        process_table(table, processed_tables)
#    print(f"📝 Result metrics['credits']: {metrics['credits']}")
#    print(f"📝 Result total_requested_credit: {metrics['total_requested_credit']}")
#    return financial_data 
#    return metrics

def extract_credit_payments_from_json(json_data):
    """Витягує дані про платежі, підтримуючи ліміти з 0 залишком (напр. Монобанк)."""
    payments_by_borrower = {}
    total_payments = 0.0
    calculated_credit_details = []
    wcf_otherbanks_total = 0.0
    if_otherbanks_total = 0.0
    
    current_date = datetime.now()
    processed_in_math = set()
    
    borrower_limits = {}

    
    def clean_val(val):
        if not val: return 0.0
        v = str(val).strip().replace(" ", "").replace(",", ".")
        try:
            return float(v) if v and v != "." else 0.0
        except:
            return 0.0
    
    for table in json_data.get("tables", []):
        table_id = table.get("table_id", "")
        
        # Обробляємо таблиці 6.2.1.11 або 6.2.11
        if table_id in ["Table_6.2.1.11", "Table_6.2.11", "6.2.1.11", "6.2.11"]:
            if table_id in processed_in_math:
                continue
            
            print(f"📝 Обробка таблиці для математики та лімітів: {table_id}")
            rows = table.get("rows", [])
            
            for row_idx, row in enumerate(rows[1:], start=2):
                cells = [str(c).strip() for c in row.get("cells", [])]                
                # Пропускаємо підсумкові рядки
                if any(cell.lower().startswith("всього") or cell.lower().startswith("загальна") for cell in cells):
                    continue
                
                if len(cells) >= 10:
                    try:
                        # 1. Отримання базових значень (додаємо перевірку на довжину та вміст)
                        limit_raw = cells[3] if len(cells) > 3 else "0"
                        balance_raw = cells[4] if len(cells) > 4 else "0"

                        limit = clean_val(limit_raw) * 1000
                        balance = clean_val(balance_raw) * 1000
                        l_type = cells[2].strip().lower() if cells[2] else "кредит"
                        calculated_credit_details.append({
                            "type": l_type,
                            "borrower": cells[1],
                            "limit": limit,
                            "balance": balance,
                            "purpose": cells[7] if len(cells) > 7 else ""
                        })
                        # Тип ліміту - якщо порожньо в cells[2], спробуємо взяти назву з цілі (cells[8])
                        limit_type = cells[2].strip().lower() if len(cells) > 2 and cells[2] else "кредит"

                        # 2. Визначення ID позичальника (якщо немає дужок, беремо весь текст або "Клієнт")
                        borrower_id_raw = cells[1].strip() if len(cells) > 1 else ""
                        bid_match = re.search(r'\((\d+)\)', borrower_id_raw)

                        if bid_match:
                            f_bid = bid_match.group(1)
                        elif borrower_id_raw:
                            f_bid = borrower_id_raw # Беремо назву текстом, якщо немає цифр в дужках
                        else:
                            f_bid = "Основний позичальник"

                        if limit > 0 or balance > 0:
                            # --- ЛОГІКА РОЗРАХУНКУ РІЧНОГО ПЛАТЕЖУ (для DSCR) ---
                            rate = clean_val(cells[5]) / 100 if cells[5] else 0.25
                            
                            # Парсинг дати та розрахунок терміну
                            maturity_date_raw = cells[9].strip().replace(" 00:00:00 0", "") if len(cells) > 9 else None
                            term_months = 12 # дефолт
                            if maturity_date_raw:
                                try:
                                    end_dt = datetime.strptime(maturity_date_raw, "%d.%m.%Y")
                                    diff = end_dt - current_date
                                    term_months = max(1, diff.days // 30)
                                except:
                                    pass

                            if any(x in limit_type for x in ["вкл", "кк", "овердрафт", "лінія", "овд"]):
                                annual_payment = limit * rate
                            else:
                                annual_principal = balance / (term_months / 12) if term_months > 12 else balance
                                annual_payment = annual_principal + (balance * rate)

                            payments_by_borrower[f_bid] = payments_by_borrower.get(f_bid, 0.0) + annual_payment
                            total_payments += annual_payment

                            # --- НОВА ЛОГІКА КЛАСИФІКАЦІЇ ЛІМІТІВ (WCF/IF) ---
                            # Вибір бази: для ліній - ліміт, для іншого - залишок
                            if any(x in limit_type for x in ["кк", "вкл", "овд", "овердрафт"]):
                                calc_sum = limit
                            else:
                                calc_sum = balance

                            # Визначення типу ліміту за ціллю (зазвичай cells[8])
                            # Якщо індекс цілі інший - підправте тут
                            purpose = cells[7].lower() if len(cells) > 7 else ""
                            
                            if "придб" in purpose:
                                if_otherbanks_total += calc_sum
                            else:
                                wcf_otherbanks_total += calc_sum

                            print(f"✅ Рядок {row_idx}: {limit_type} | Сума для ліміту: {calc_sum:,.2f} | Платіж: {annual_payment:,.2f}")

                    except Exception as e:
                        print(f"⚠️ Помилка в рядку {row_idx}: {e}")
            
            processed_in_math.add(table_id)
            break 
    
    return {
        "total_annual_payments_other_banks": total_payments,
        "payments_by_borrower": payments_by_borrower,
        "WCF_otherbanks_total": wcf_otherbanks_total,
        "IF_otherbanks_total": if_otherbanks_total,
        "calculated_credit_details": calculated_credit_details
    }

import re
from datetime import datetime
from dateutil.relativedelta import relativedelta

def extract_credit_history_payments(json_data):
    # Зберігаємо всі ваші оригінальні назви
    payments_by_borrower = {}
    total_payments = 0.0
    credit_history_payments = []
    processed_table_ids = set()
    current_date = datetime.now()
    
    wcf_rbua_total = 0.0    
    if_rbua_total = 0.0

    def clean_val(val):
        if not val: return "0"
        v = str(val).strip().replace(" ", "").replace("\xa0", "").replace(",", ".")
        return v if v not in ["", ".", ".00"] else "0"

    def parse_date(d_str):
        if not d_str or not isinstance(d_str, str) or len(d_str) < 5: return None
        d_clean = d_str.strip().split(" ")[0]
        for fmt in ("%Y-%m-%d", "%d.%m.%Y"):
            try:
                return datetime.strptime(d_clean, fmt)
            except: continue
        return None

    all_tables = []
    for t in json_data.get("tables", []):
        all_tables.append(t)
        for row in t.get("rows", []):
            for nt in row.get("nested_tables", []):
                all_tables.append(nt)

    for table in all_tables:
        tid = str(table.get("table_id", ""))
        title = str(table.get("title", ""))
        
        is_target_table = (
            any(x in tid for x in ["5.2.11", "5.2.1.11", "Table_5"]) or 
            "кредитн" in title.lower() or 
            "діючи" in title.lower()
        )

        if is_target_table:
            if tid in processed_table_ids: 
                continue
            processed_table_ids.add(tid)
            
            rows = table.get("rows", [])
            for row in rows:
                cells = row.get("cells", [])
                
                if not cells: continue
                first_cell = str(cells[0]).strip()
                
                # Залишаємо фільтр тільки для технічних заголовків
                if first_cell in ["Банк", "Назва банку"]:
                    continue

                try:
                    limit = float(clean_val(cells[4])) if len(cells) > 4 else 0.0
                    balance = float(clean_val(cells[6])) if len(cells) > 6 else 0.0
                    
                    if limit >= 0 or balance >= 0:
                        final_bid = cells[1].strip() if len(cells) > 1 and cells[1].strip() else cells[0].strip()
                        
                        # ВИЗНАЧАЄМО ЧИ ЦЕ ПІДСУМКОВИЙ РЯДОК
                        is_total = any(word in final_bid.lower() for word in ["всього", "разом", "загальна", "підсумок"])

                        if is_total:
                            # Для підсумку ставимо прочерки і НЕ рахуємо математику
                            s_date = "-"
                            e_date = "-"
                            display_rate = "-"
                        else:
                            # Для звичайного кредиту парсимо дати та рахуємо показники
                            issue_dt = parse_date(cells[2]) or (current_date - relativedelta(years=1))
                            end_dt = parse_date(cells[3]) or (current_date + relativedelta(years=1))
                            s_date = issue_dt.strftime("%d.%m.%Y")
                            e_date = end_dt.strftime("%d.%m.%Y")
                            
                            rate_raw = clean_val(cells[7]) if len(cells) > 7 else "25"
                            rate_val = float(rate_raw) / 100 if float(rate_raw) > 0 else 0.25
                            display_rate = f"{rate_val*100:.1f}%"

                            # РОЗРАХУНОК (тільки для не-підсумкових рядків)
                            delta = relativedelta(end_dt, issue_dt)
                            total_term_months = max(1, delta.years * 12 + delta.months)
                            
                            # Парсимо дати
                            issue_dt = parse_date(cells[2]) or (current_date - relativedelta(years=1))
                            end_dt = parse_date(cells[3]) or (current_date + relativedelta(years=1))
                            
                            s_date = issue_dt.strftime("%d.%m.%Y")
                            e_date = end_dt.strftime("%d.%m.%Y")
                            
                            rate_raw = clean_val(cells[7]) if len(cells) > 7 else "25"
                            rate_val = float(rate_raw) / 100 if float(rate_raw) > 0 else 0.25
                            display_rate = f"{rate_val*100:.1f}%"

                            # ПЕРЕВІРКА: чи кредит ще діючий?
                            # Додаємо в розрахунки тільки якщо дата закінчення >= сьогодні
                            if end_dt >= current_date:
                                delta = relativedelta(end_dt, issue_dt)
                                total_term_months = max(1, delta.years * 12 + delta.months)
                                
                                if balance <= 1.0:
                                    annual_payment = limit * rate_val
                                    calc_sum = limit
                                else:
                                    annual_principal = balance / (total_term_months / 12) if total_term_months > 0 else balance
                                    annual_payment = annual_principal + (balance * rate_val)
                                    calc_sum = balance

                                category = "IF_RBUA" if total_term_months > 48 else "WCF_RBUA"
                                if category == "IF_RBUA": if_rbua_total += calc_sum
                                else: wcf_rbua_total += calc_sum
                                
                                total_payments += annual_payment
                                payments_by_borrower[final_bid] = payments_by_borrower.get(final_bid, 0.0) + annual_payment
                            else:
                                # Якщо кредит прострочений або закритий, 
                                # ми НЕ додаємо його в суми, але він залишиться в таблиці нижче
                                pass

                        # Додаємо в масив для таблиці (і звичайні, і підсумки)
                        credit_history_payments.append({
                            "borrower": final_bid,
                            "limit": limit,
                            "balance": balance,
                            "start_date": s_date,
                            "end_date": e_date,
                            "rate": display_rate
                        })
                except: continue

    return {
        "total_annual_payments_history": total_payments,
        "WCF_RBUA_total": wcf_rbua_total,
        "IF_RBUA_total": if_rbua_total,
        "credit_history_payments": credit_history_payments,
        "payments_by_borrower": payments_by_borrower,
        "final_risk_manager": json_data.get("risk_manager") or "Уніченко Лілія"
    }
def create_new_credit_history_table(doc, credit_history_payments):

    if not credit_history_payments:
        print("⚠️ Дані кредитної історії відсутні.")
        return
     
    # 1. Спочатку визначаємо заголовки
    headers = ['Банк', 'Позичальник', 'Дата видачі', 'Дата погашення', 'Ліміт', 'Валюта', 'Залишок', 'Ставка']
    
    # 2. Тепер створюємо таблицю (кількість колонок = довжині списку заголовків)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)

    # 3. Заповнюємо шапку
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        # Опціонально: робимо текст жирним
        if hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs[0].bold = True
    history_list = []
    if isinstance(credit_history_payments, dict):
        history_list = credit_history_payments.get('credit_history_payments', [])
    else:
        history_list = credit_history_payments
    # 4. Заповнюємо дані
    for item in history_list:
        row_cells = table.add_row().cells
        
        row_cells[0].text = "АТ 'РАЙФФАЙЗЕН БАНК'"
        row_cells[1].text = str(item.get('borrower', ''))
        row_cells[2].text = str(item.get('start_date', ''))
        row_cells[3].text = str(item.get('end_date', ''))
        
        limit = item.get('limit', 0.0)
        row_cells[4].text = f"{limit:,.2f}".replace(",", " ")
        
        row_cells[5].text = "UAH"
        
        balance = item.get('balance', 0.0)
        row_cells[6].text = f"{balance:,.2f}".replace(",", " ")
        
        row_cells[7].text = str(item.get('rate', ''))
   
def load_doc_text(uploaded_file):
    """Load and parse DOCX file into text and JSON structure."""
    try:
        doc = Document(uploaded_file)
        full_text = []
        json_output = {
            "paragraphs": [],
            "tables": [],
            # Поля залишаємо для сумісності, вони заповняться пізніше іншим кодом
            "deal_description": "Дані про угоду відсутні.",
            "ownership_info": "Інформація про власників відсутня.",
            "business_description": "Інформація про бізнес відсутня.",
            "mtb_info": "Дані про МТБ відсутні.",
            "war_impact": "Дані про вплив воєнного стану відсутні.",
            "revenue_change_analysis": "Дані відсутні.",
            "capital_withdrawal": "Дані відсутні.",
            "profit_usage": "Дані відсутні.",
            "main_contractors": "Дані відсутні."
        }
        print(f"📝 Processing document: {uploaded_file}")
        
        # 1. Обробка параграфів
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                normalized_text = unicodedata.normalize("NFC", text).replace('\xa0', ' ').replace('’', "'").replace('\u2019', "'").replace('\u2018', "'")
                full_text.append(normalized_text)
                json_output["paragraphs"].append(normalized_text)
                # print(f"📝 Added paragraph: {normalized_text[:50]}...") # Опціонально для дебагу
        
        # 2. Рекурсивна функція для таблиць
        def process_table(table, table_idx, parent_path=""):
            table_id = f"{parent_path}{table_idx}" if parent_path else f"Table_{table_idx}"
            table_data = {"table_id": table_id, "title": "", "rows": []}
            
            if table.rows and table.rows[0].cells:
                first_row_text = " ".join(cell.text.strip() for cell in table.rows[0].cells if cell.text.strip())
                table_data["title"] = first_row_text
                
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                nested_tables = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = unicodedata.normalize("NFC", cell.text.strip()).replace('\xa0', ' ').replace('’', "'").replace('\u2019', "'").replace('\u2018', "'")
                    row_cells.append(cell_text if cell_text else "")
                    
                    # Обробка вкладених таблиць
                    for sub_table_idx, sub_table in enumerate(cell.tables):
                        sub_table_path = f"{table_id}.{row_idx + 1}.{cell_idx + 1}.{sub_table_idx + 1}"
                        nested_tables.append(process_table(sub_table, sub_table_idx + 1, sub_table_path))
                
                row_data = {"cells": row_cells, "nested_tables": nested_tables, "row_idx": row_idx + 1}
                if any(row_cells) or nested_tables:
                    table_data["rows"].append(row_data)
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        full_text.append(row_text)
            
            json_output["tables"].append(table_data)
            return table_data
        
        # 3. Запуск обробки таблиць
        for table_idx, table in enumerate(doc.tables):
            process_table(table, table_idx + 1)
        
        # 4. Збереження JSON (як у вашому коді)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        results_dir = "results"
        os.makedirs(results_dir, exist_ok=True)
        json_filename = f"Document_Structure_{timestamp}.json"
        json_filepath = os.path.join(results_dir, json_filename)
        
        with open(json_filepath, "w", encoding="utf-8") as f:
            json.dump(json_output, f, ensure_ascii=False, indent=2)
            print(f"📝 Saved JSON structure: {json_filepath}")
        
        if not full_text:
            print(f"⚠️ No text extracted from document: {uploaded_file}")
        else:
            print(f"✅ Successfully processed document, length: {len(full_text)}")
            
        return "\n".join(full_text), json_output

    except Exception as e:
        print(f"⚠️ Error processing document {uploaded_file}: {e}")
        return "", {}

def evaluate_company(text, text_rules=None, financial_data=None, financial_signals=None):
    import re
    from datetime import datetime
    
    # 0. Ініціалізація змінних (щоб уникнути UnboundLocalError)
    result = {"Плюси": [], "Мінуси": []}
    contr_data = [] # Тепер вона точно існує з самого початку
    
    if text_rules is None: text_rules = {}
    if financial_signals is None: financial_signals = []

    # 1. Текстовий аналіз
    for section in ["Плюси", "Мінуси", "Ризики", "+/-"]:
        rules_list = text_rules.get(section) or []
        for rule in rules_list:
            try:
                pattern = rule.get("pattern")
                if pattern and re.search(pattern, text, re.IGNORECASE):
                    comment = rule.get("comment")
                    if not comment: continue
                    target = "Мінуси" if section == "Ризики" else section
                    target = "Плюси" if section == "+/-" else target
                    result[target].append(f"{'+' if target == 'Плюси' else '-'} {comment}")
            except Exception: continue

    # 2. Фінансові сигнали
    for signal in financial_signals:
        comment = signal.get("comment", "")
        target = "Мінуси" if any(w in comment.lower() for w in ["відхилення", "низька"]) else "Плюси"
        result[target].append(f"{'+' if target == 'Плюси' else '-'} {comment}")

    # 3. Визначення року (Крок А і Б)
    start_year = None
    pattern_reg = r"(?:зареєстр|діє з|працює з|ФОП з|із |з |заснован|реєстрації|початок діяльності).{1,20}?\b((?:19|20)\d{2})\b"
    match_reg = re.search(pattern_reg, text, re.IGNORECASE | re.DOTALL)
    
    if match_reg:
        start_year = int(match_reg.group(1))
        print(f"🎯 ЗНАЙДЕНО РІК ЗА МАРКЕРОМ: {start_year}")
    
    if not start_year:
        all_years = re.findall(r"\b(19\d{2}|20[0-2]\d)\b", text)
        if all_years:
            start_year = min(int(y) for y in all_years if 1990 <= int(y) <= datetime.now().year)

    # --- 4. ПОДАЛЬША ЛОГІКА (Тепер вона ПОЗА блоком 'if not start_year') ---
    if start_year:
        experience = max(datetime.now().year - start_year, 1)
        print(f"📊 ПІДСУМОК: Рік={start_year}, Досвід={experience} р.")

        if experience <= 2:
            result["Мінуси"].append(f"- Незначний досвід діяльності ({experience} р.)")
        else:
            msg = "+ Значний досвід діяльності" if experience <= 5 else f"+ Має досвід діяльності більше 5 років ({experience} р.)"
            result["Плюси"].append(msg)

        if financial_data:
            try:
                ebitda = float(financial_data.get("ebitda", 0) or 0)
                equity = float(financial_data.get("equity", 0) or 0)
                calc_exp = min(experience, 10)
                expected_inv = calc_exp * ebitda * 0.20
                print(f"💰 ДІАГНОСТИКА: EBITDA={ebitda}, Equity={equity}, Exp_Inv={expected_inv}")

                if equity <= 0:
                    result["Мінуси"].append("- Відсутній власний капітал (негативна капіталізація)")
                elif expected_inv <= equity:
                    result["Плюси"].append("+ Клієнт вкладає кошти в бізнес (рівень капіталу відповідає терміну діяльності)")
                else:
                    result["Мінуси"].append("- Значне виведення коштів на власні цілі")
            except: pass
    else:
        result["Мінуси"].append("- Неможливо оцінити досвід діяльності")

    return result

def get_calculation_params_from_json(json_data):
    tables = json_data.get("tables", [])
    for table in tables:
        if table.get("table_id") == "Table_14":
            extracted_rows = []
            for row in table.get("rows", []):
                cells = row.get("cells", [])
                # Беремо всі колонки, які є в рядку
                extracted_rows.append(cells)
            return extracted_rows
    return None

def save_results_to_docx(financial_signals, analyst_name, credit_data_res, result, timestamp, credit_payments_info, credit_history_payments, json_data, financial_data, credit_payments, calc_params_table, descriptions, all_paragraphs, doc):
    """Generate and save the underwriter's conclusion as a DOCX file."""
    credit_pattern = r"\b(\d{1,3}(?:\s\d{3})*,\d{0,2})\s*(?:грн\.?|UAH)?\b"
    
        # Set custom page margins (1 cm = 20/20 inch)
    section = doc.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Combine full_text from load_doc_text for better data access
    full_text = '\n'.join(json_data.get('paragraphs', []) + [row_text for table in json_data.get('tables', []) for row in table.get('rows', []) for row_text in [' | '.join(row.get('cells', []))]])
    print("🔍 Diagnostics: json_data keys:", list(json_data.keys()))
    print("🔍 Diagnostics: financial_data keys:", list(financial_data.keys()))
    print("🔍 Diagnostics: full_text sample:", full_text[:200])  # Log first 200 chars for debugging
    print("🔍 Diagnostics: tables:", [t.get('title', t.get('table_id', 'No ID')) for t in json_data.get('tables', [])])  # Log table titles/ids
    
    # Document header
    heading = doc.add_heading(f"Висновок Андерайтера від {datetime.now().strftime('%d.%m.%Y')}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].bold = True
    
    # Bank and OD info
    p = doc.add_paragraph("Райффайзен банк")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Dynamic extraction for OD
    od_match = next((line for line in full_text.split('\n') if re.search(r'(?:Дніпропетровська|Одеська|Чернігівська|Черкаська|Харківська|Полтавська|Сумська|Львівська|Кіровоградська|Миколаївська|Одеська|Запорізька|Івано-Франківська|Хмельницька|Тернопільська|Чернівецька|Волинська|Рівненська|Ужгородська)\s+ОД', line, re.IGNORECASE)), "Київська РД")
    p = doc.add_paragraph(od_match)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Dynamic extraction for Department
    dept_match = next((line for line in full_text.split('\n') if re.search(r'Відділення\s+[\'"]?[^\s]+[\'"]?', line, re.IGNORECASE)), "Невідоме відділення")
    p = doc.add_paragraph(dept_match)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Manager and client info
    client_manager = json_data.get('client_manager', financial_data.get('manager_name', None))
    risk_manager = json_data.get('risk_manager', financial_data.get('underwriter_name', None))
    if not client_manager:
        # Search in full_text
        client_manager_match = re.search(r'(?:Менеджер\s*(?:клієнта|:)|\s*Manager\s*(?:client|:)|\s*Менеджер)\s*[:\s]*([^\n]+)', full_text, re.IGNORECASE)
        client_manager = client_manager_match.group(1).strip() if client_manager_match else None
        if not client_manager:
            for table in json_data.get('tables', []):
                for row in table.get('rows', []):
                    cells = row.get('cells', [])
                    for i, cell in enumerate(cells):
                        if re.search(r'(?:Менеджер\s*(?:клієнта|)|Manager\s*(?:client|)|Менеджер)', cell, re.IGNORECASE) and i + 1 < len(cells):
                            client_manager = cells[i + 1].strip()
        client_manager = client_manager or 'Невідомий менеджер'
        
        if analyst_name and analyst_name.strip():
            final_risk_manager = analyst_name
        else:
            # Беремо значення з JSON або дефолтне, якщо JSON порожній
            final_risk_manager = json_data.get("risk_manager") or "Уніченко Лілія"
    if client_manager == 'Невідомий менеджер' or final_risk_manager == 'Невідомий ризик-менеджер':
            print(f"⚠️ Manager or risk manager names not found. Check full_text or tables for variations of 'Менеджер' or 'Ризик менеджер'.")
    
    p = doc.add_paragraph(f"Менеджер клієнта: {client_manager}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"Ризик менеджер: {final_risk_manager}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Client info
    client_name = financial_data.get('borrower_status', next((line.split('/')[0].strip() for line in full_text.split('\n') if '/' in line), 'Невідомий клієнт'))
    client_id = financial_data.get('kved_code', next((line.split('/')[1].strip() for line in full_text.split('\n') if '/' in line), 'Невідомий код'))
    p = doc.add_paragraph(f"{client_name} / {client_id}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Dynamic extraction for addresses and registration date
    address_location = next((line.replace('Адреса за місцем реєстрації', '').strip() for line in full_text.split('\n') if 'Адреса за місцем реєстрації' in line), 'Невідома адреса')
    address_work = next((line.replace('Адреса за місцем роботи', '').strip() for line in full_text.split('\n') if 'Адреса за місцем роботи' in line), 'Невідома адреса')
    reg_date = next((line.replace('Дата реєстрації:', '').strip() for line in full_text.split('\n') if 'Дата реєстрації:' in line), 'Невідома дата')
    p = doc.add_paragraph(f"Адреса місця знаходження клієнта: {address_location}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"Адреса за місцем роботи: {address_work}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"Дата реєстрації: {reg_date}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")
        
    # Credit requests in table format
    doc.add_heading("Запити на кредит", level=1)
    credit_requests = financial_data.get('credits', [])
    target_table = "Table_1.11.1.11"
    if not credit_requests:
        for table in [t for t in json_data.get('tables', []) if t.get('table_id') == "Table_1.11.1.11"]:
            for row_idx, row in enumerate(table.get('rows', []), start=1):
                cells = row.get('cells', [])
                if len(cells) >= 5 and row_idx == 2:  # Second row of credit table
                    amount_str = cells[0].strip()
                    if re.match(credit_pattern, amount_str):
                        cleaned_amount = amount_str.replace(" ", "").replace(",", ".").replace("грн", "").strip()
                        amount = float(cleaned_amount) * 1000
                        term = cells[2].strip() if cells[2].strip().isdigit() else "36"
                        rate = cells[3].strip().replace("%", "") if cells[3].strip() else "20"
                        try:
                            rate = float(rate) / 100
                        except ValueError:
                            rate = 0.20
                        product = cells[4].strip()
                        credit_requests.append({
                            "amount": amount,
                            "currency": "UAH",  # Assuming UAH as default currency
                            "term": term,
                            "rate": rate,
                            "type": product if product else "Невідомий"
                        })
                        print(f"✅ Extracted credit: {amount} грн, term: {term}, rate: {rate}, type: {product}")
    
    # Create a table for each credit request
    for request in credit_requests:
        credit_table = doc.add_table(rows=3, cols=5)
        set_table_borders(credit_table)
        credit_table.autofit = False
        
        # Headers
        headers = ["Сума", "Валюта", "Термін", "Ставка", "Стандарт"]
        for i, header in enumerate(headers):
            cell = credit_table.rows[0].cells[i]
            # Очищуємо комірку і додаємо новий run
            cell.text = "" 
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(3.0)
        
        # Data row
        row_cells = credit_table.rows[1].cells
        row_cells[0].text = "{:,.0f}".format(request.get('amount', 0) / 1000).replace(',', ' ')  # e.g., "29 000"
        row_cells[1].text = request.get('currency', 'UAH')
        row_cells[2].text = str(request.get('term', 'Невідомий'))
        row_cells[3].text = f"{request.get('rate', 0.20) * 100:.2f}%"
        row_cells[4].text = request.get('type', 'Невідомий')
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(3.0)
        
        # Purpose row as a single merged cell
        purpose = "Обігові кошти" if any(kw in request.get('type', '').lower() for kw in ["довіра", "овердрафт", "вкл"]) else \
                    "Розвиток бізнесу" if "розвиток" in request.get('type', '').lower() else \
                    "Інвестиційні цілі" if "інвест" in request.get('type', '').lower() else "Цільове призначення не визначено"
    
    doc.add_paragraph("")
    
    # Verification table
    doc.add_heading("Верифікація", level=1)
    verification_table = doc.add_table(rows=7, cols=2)  # 7 rows as specified
    set_table_borders(verification_table)
    verification_table.autofit = False
    
    headers = ["Параметр", "Коментарі"]
    for i, header in enumerate(headers):
        cell = verification_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Cm(8.0 if i == 0 else 11.0)  # Adjust widths for better distribution
    
    verification_data = [
        "Клієнт та ГПК у зовнішній БД",
        "Warning list (клієнт та ГПК)",
        "Банкрутство (клієнт та ГПК)",
        "КБ",
        "Реєстр судових рішень",
        "Висновок СБ",
        ""  # Empty row as specified
    ]
    
    for i in range(7):  # Iterate over 0 to 6 to match 7 rows
        row_cells = verification_table.rows[i].cells
        if i == 0:
            continue  # Skip header row
        param = verification_data[i - 1] if i - 1 < len(verification_data) else ""
        row_cells[0].text = param
        row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[0].width = Cm(8.0)
        row_cells[1].text = ""  # Placeholder for comments
        row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].width = Cm(11.0)
    
    doc.add_paragraph("")
    
    # Auto-check results table
    auto_check_table = None
    for table in json_data.get('tables', []):
        if table.get('title') and re.search(r'Результаты\s+автопроверок\s+по\s+заявке\s+№\d+/\d+/\d+', table.get('title', '')):
            auto_check_table = table
            break
        # Fallback: Check first row content if title is missing
        elif table.get('rows') and table.get('rows')[0].get('cells') and re.search(r'Результаты\s+автопроверок\s+по\s+заявке\s+№\d+/\d+/\d+', table.get('rows')[0].get('cells')[0], re.IGNORECASE):
            auto_check_table = table
            table['title'] = table['rows'][0]['cells'][0].strip()  # Set title if found in first cell
            break
    if auto_check_table:
        rows = auto_check_table.get('rows', [])
        if len(rows) >= 2:  # Ensure at least header and data rows
            # Extract request number for title
            request_number = re.search(r'№(\d+/\d+/\d+)', auto_check_table.get('title', '')).group(1) if re.search(r'№\d+/\d+/\d+', auto_check_table.get('title', '')) else "Невідомий номер"
            doc.add_heading(f"Результати автоперевірок по заявці №{request_number}", level=1)
            
            # Use second row as headers, translate to Ukrainian
            headers = []
            if len(rows) > 1 and rows[1].get('cells'):
                raw_headers = [cell.strip() for cell in rows[1].get('cells', [])[:20]]  # Limit to 20 columns
                translation_map = {
                    "Тип стороны": "Тип сторони",
                    "ИНН/ОКПО": "ІПН/ЄДРПОУ",
                    "ФИО/Название": "ПІБ/Назва",
                    "Сегмент": "Сегмент",
                    "Количество действующих кредитов": "Кількість діючих кредитів",
                    "Действующих с нулевым балансом": "Діючих із нульовим балансом",
                    "Количество погашенных кредитов": "Кількість погашених кредитів",
                    "Утвержденный лимит по договорам, тыс грн": "Затверджений ліміт за договорами, тис грн",
                    "Расчетные платежи по кредитам, тыс. грн.": "Розрахункові платежі за кредитами, тис грн",
                    "Макс. DPD за 12 мес": "Макс. DPD за 12 міс",
                    "Макс. DPD ever": "Макс. DPD ever",
                    "DPD на сегодня": "DPD на сьогодні",
                    "PD": "PD",
                    "Блек-лист": "Warning list",
                    "Инсайдер": "Інсайдер",
                    "Работник банка": "Працівник банку",
                    "С/м поступления тыс. грн.": "Середньомісячні надходження, тис грн",
                    "Срок пользования счетом": "Термін користування рахунком",
                    "Дата рождения": "Дата народження",
                    "Возраст": "Вік"
                }
                headers = [translation_map.get(h, h) for h in raw_headers][:20]  # Translate and limit to 20
            
            # Prepare unique rows based on ІПН/ЄДРПОУ
            unique_rows = {}
            for row in rows[2:]:  # Skip title and header rows
                cells = row.get('cells', [])
                if len(cells) >= 2:  # Ensure at least ІПН/ЄДРПОУ and other data
                    inn_okpo = cells[1].strip()  # Column 1 is ІПН/ЄДРПОУ
                    if inn_okpo and inn_okpo not in unique_rows:
                        unique_rows[inn_okpo] = cells[:20]  # Take up to 20 columns
            
            # Create table with dynamic rows
            auto_table = doc.add_table(rows=len(unique_rows) + 1, cols=20)  # +1 for header
            set_table_borders(auto_table)
            auto_table.autofit = False
            
            # Headers
            for i, header in enumerate(headers):
                cell = auto_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.size = Pt(6)  # Set font size to 6
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if i + 1 in [2, 3, 4]:  # Columns 2, 3, 4
                    cell.width = Cm(1.5)  # Fixed width for columns 2, 3, 4
                else:  # Columns 1, 5-20
                    cell.width = Cm(0.8)  # Adjusted for remaining 17 columns 
            
            # Data rows with color handling
            from docx.oxml.ns import qn

            for idx, (inn_okpo, cells) in enumerate(unique_rows.items(), start=1):
                row_cells = auto_table.rows[idx].cells
                for i, value in enumerate(cells[:20]):  # Limit to 20 columns
                    row_cells[i].text = str(value) if i < len(cells) else ""
                    row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(6)  # Set font size to 6
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if i + 1 in [2, 3, 4]:  # Columns 2, 3, 4
                        row_cells[i].width = Cm(1.5)  # Fixed width for columns 2, 3, 4
                    else:  # Columns 1, 5-20
                        row_cells[i].width = Cm(0.8)  # Adjusted for remaining 17 columns
                    # Color logic based on requirements
                    tc = row_cells[i]._tc  # Get the underlying XML element
                    tcPr = tc.get_or_add_tcPr()  # Get or add tcPr element
                    shd = tcPr.first_child_found_in("w:shd")  # Get existing shd, or None
                    if shd is None:
                        shd = OxmlElement('w:shd')  # Create new shd element if it doesn't exist
                        tcPr.append(shd)
                    shd.set(qn('w:fill'), 'auto')  # Default to 'auto' to clear previous fill
                    if i == 0:  # Column 1: Тип сторони
                        if str(value).upper() == "CLN":
                            shd.set(qn('w:fill'), '92D050')  # Green fill
                        elif str(value).upper() == "FND":
                            shd.set(qn('w:fill'), 'FFFF00')  # Yellow fill
                        elif str(value).upper() == "LNK":
                            shd.set(qn('w:fill'), 'FFA500')  # Orange fill
                    elif i in [9, 10]:  # Columns 9, 10: Макс. DPD за 12 міс, Макс. DPD ever
                        try:
                            dpd_value = float(str(value).replace(',', '.').replace('%', ''))
                            if dpd_value != 0 and str(value).strip():
                                shd.set(qn('w:fill'), 'FFFF00')  # Yellow fill
                        except ValueError:
                            pass
                    elif i in [13, 14, 15]:  # Columns 13, 14, 15: Warning list, Інсайдер, Працівник банку
                        try:
                            if float(str(value).replace(',', '.')) != 0:
                                shd.set(qn('w:fill'), 'FF0000')  # Red fill
                        except ValueError:
                            pass
                    elif i == 19:  # Column 19: Вік
                        try:
                            age = float(str(value).replace(',', '.'))
                            if age > 70:
                                shd.set(qn('w:fill'), 'FF0000')  # Red fill
                        except ValueError:
                            pass
            print(f"📝 Added auto-check table with {len(unique_rows)} unique rows")
    else:
        print("⚠️ Auto-check table not found in json_data['tables']. Check if title or first row content matches the expected pattern.")
    
    doc.add_paragraph("")

    # Calculated parameters
    doc.add_heading("Розрахункові параметри", level=1)
    params_table = doc.add_table(rows=1, cols=3)
    set_table_borders(params_table)
    params_table.autofit = True
    headers = ["Найменування параметра", "Значення параметра", "Повідомлення"]
    for i, header in enumerate(headers):
        cell = params_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Cm(4.0)
    
    param_rows = [
        ["02. Назва/ПІБ та ІПН/ЄДРПОУ клієнта", f"{client_name}, ІПН: {client_id}", ""],
        ["03. Дата реєстрації", reg_date, ""],
        ["04. Адреса реєстрації", address_location, ""],
    ]
    for row in param_rows:
        row_cells = params_table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[i].width = Cm(4.0)
       
    doc.add_paragraph("")
    p_title = doc.add_paragraph("Кредитна історія")
    p_title.runs[0].bold = True
    
    # ПРАВИЛЬНА ПЕРЕВІРКА: дістаємо список із пакета
    items_to_process = []
    if isinstance(credit_history_payments, dict):
        items_to_process = credit_history_payments.get('credit_history_payments', [])
    else:
        items_to_process = credit_history_payments
    print(f"DEBUG: Кількість елементів у кредитній історії: {len(items_to_process)}")

    if items_to_process:
        
        headers = ["Банк", "Позичальник", "Видача", "Погашення", "Ліміт", "Валюта", "Залишок", "Ставка"]
        hist_table = doc.add_table(rows=1, cols=len(headers))
        set_table_borders(hist_table)
        hist_table.autofit = True
        
        # Заголовки
        h_cells = hist_table.rows[0].cells
        for i, text in enumerate(headers):
            h_cells[i].text = text
            if h_cells[i].paragraphs[0].runs:
                h_cells[i].paragraphs[0].runs[0].bold = True
                h_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            h_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Наповнення даними
        for item in items_to_process:
            if not isinstance(item, dict): continue
            
            row_cells = hist_table.add_row().cells
            row_cells[0].text = "АТ 'РАЙФФАЙЗЕН БАНК'"
            row_cells[1].text = str(item.get('borrower', ''))
            row_cells[2].text = str(item.get('start_date', ''))
            row_cells[3].text = str(item.get('end_date', ''))
            
            limit = item.get('limit', 0.0)
            row_cells[4].text = f"{limit:,.2f}".replace(",", " ")
            row_cells[5].text = "UAH"
            
            balance = item.get('balance', 0.0)
            row_cells[6].text = f"{balance:,.2f}".replace(",", " ")
            row_cells[7].text = str(item.get('rate', ''))
    else:
        doc.add_paragraph("Дані про кредитну історію відсутні.")
    # Credit history table (conditional)
    
    # Діючі кредити клієнта та членів ГПК в інших банках

    doc.add_paragraph("")  # відступ перед таблицею

    p = doc.add_paragraph("Діючі кредити клієнта та членів ГПК в інших банках:")
    p.runs[0].bold = True
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)

    other_bank_data = credit_history_payments.get("other_bank_credits", [])

    if not other_bank_data:
        print("Diagnostics: No other_bank_credits in credit_history_payments, searching json_data")
        for table in json_data.get("tables", []):
            title = table.get("title", "")
            if title.startswith("Результаты автопроверок") or title.startswith("Кредитна історія"):
                continue

            header_found = False
            temp_data = []
            for row in table.get("rows", []):
                cells = row.get("cells", [])
                if len(cells) >= 13:
                    # Перший рядок з 13+ колонками — вважаємо заголовком
                    if not header_found:
                        header_found = True
                        continue  # пропускаємо заголовок
                    temp_data.append(cells[:13])

            if temp_data:
                other_bank_data = temp_data
                break

    # === Тепер точно перевіряємо, чи є дані ===
    if not other_bank_data:
        p = doc.add_paragraph("Відсутні діючі кредити в інших банках")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        print("Diagnostics: No other-bank credits found – added placeholder")
    else:
        # Створюємо таблицю: 1 рядок заголовка + кількість рядків даних
        total_rows = 1 + len(other_bank_data)
        other_bank_table = doc.add_table(rows=total_rows, cols=13)
        set_table_borders(other_bank_table)
        other_bank_table.autofit = False

        # --- Заголовки ---
        headers = [
            "Банк", "Позичальник (ЄДРПОУ)", "Тип ліміту", "Ліміт", "Залишок",
            "% ставка", "Вид Забезпечення", "Цільове використання",
            "Дата видачі", "Дата погашення", "Валюта", "Графік погашення",
            "Щорічний платіж"
        ]
        for j, hdr in enumerate(headers):
            cell = other_bank_table.rows[0].cells[j]
            cell.text = hdr
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(2.0) if j in [0, 1, 2, 7, 8, 9, 10, 11, 12] else Cm(1.5)

        # --- Дані ---
        for i, row_data in enumerate(other_bank_data, start=1):
            for j, value in enumerate(row_data[:13]):
                cell = other_bank_table.rows[i].cells[j]
                cell.text = str(value) if value else ""
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.width = Cm(2.0) if j in [0, 1, 2, 7, 8, 9, 10, 11, 12] else Cm(1.5)

        print(f"Diagnostics: Added other-bank credits table with {len(other_bank_data)} data rows")
        doc.add_paragraph("")  # відступ після таблиці       

    
    # Members of the Local GCC / Ownership / Shareholders

    doc.add_paragraph("")
    
    p = doc.add_paragraph("Members of the Local GCC / Ownership / Shareholders")
    p.runs[0].bold = True
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)


    gcc_data = []
    found_gcc_table = False

    for table in json_data.get("tables", []):
        rows = table.get("rows", [])
        if not rows:
            continue

        first_row_cells = rows[0].get("cells", [])
        if len(first_row_cells) < 5:
            continue

        # Перевірка: чи є в першому рядку потрібні ключові слова
        first_row_text = " ".join(str(c).lower() for c in first_row_cells)
        if any(kw in first_row_text for kw in ["назва клієнта", "члена гпк", "% в уф", "окпо", "іпн", "urg", "власник", "тип зв’язку"]):
            found_gcc_table = True
            # Пропускаємо перший рядок (заголовки), беремо решту
            for row in rows[1:]:
                cells = row.get("cells", [])
                if len(cells) >= 5:
                    gcc_data.append(cells[:5])
            break

    if not gcc_data:
        p = doc.add_paragraph("Дані про власників/членів ГПК відсутні")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        print("Diagnostics: No GCC/Ownership data found – added placeholder")
    else:
        gcc_table = doc.add_table(rows=len(gcc_data) + 1, cols=5)
        gcc_table.autofit = False

        headers = ["Назва клієнта/члена ГПК", "% в УФ", "Назва власника/тип зв’язку", "ОКПО/ІПН", "URG"]
        for j, hdr in enumerate(headers):
            cell = gcc_table.rows[0].cells[j]
            cell.text = hdr
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(3.0) if j in [0, 2] else Cm(1.8)

        for i, row in enumerate(gcc_data, start=1):
            for j, value in enumerate(row):
                cell = gcc_table.rows[i].cells[j]
                cell.text = str(value) if value else ""
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.width = Cm(3.0) if j in [0, 2] else Cm(1.8)

        print(f"Diagnostics: Added GCC/Ownership table with {len(gcc_data)} rows")
        doc.add_paragraph("")


    # ----------------------------------------------------------------------
    # Обороти по рахункам
    # ----------------------------------------------------------------------
    doc.add_paragraph("")
    p = doc.add_paragraph("Обороти по рахункам")
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    total_inflow = 0.0
    turnover_data = []
    clean_turnover = []

    # === 1. Отримуємо revenue (виручку) ===
    revenue = financial_data.get("revenue", 0)
    if revenue:
        try:
            revenue = float(str(revenue).replace(" ", "").replace(",", "."))
            print(f"Diagnostics: Revenue taken from financial_data: {revenue} тис. грн")
        except (ValueError, TypeError):
            revenue = 0

    # === 2. Якщо виручки немає в ключі — шукаємо в тексті ===
    if revenue == 0:
        full_text = "\n".join(json_data.get("paragraphs", [])) + "\n" + \
                    "\n".join(" | ".join(row.get("cells", [])) for table in json_data.get("tables", []) for row in table.get("rows", []))

        patterns = [
            r"виручка.*?([\d\s.,]+)\s*тис",
            r"виручка.*?([\d\s.,]+)\s*грн",
            r"виручка.*?([\d\s.,]+)",
            r"всього\s+виручки.*?([\d\s.,]+)",
            r"дохід.*?([\d\s.,]+)",
            r"реалізації.*?([\d\s.,]+)",
        ]

        for pattern in patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                try:
                    val = match.group(1).replace(" ", "").replace(",", ".")
                    revenue = float(val)
                    print(f"Diagnostics: Revenue found in text: {revenue} тис. грн (pattern: {pattern})")
                    break
                except:
                    continue

    # === 3. Пошук таблиці оборотів (ігноруємо 2 рядки заголовка) ===
    for table in json_data.get("tables", []):
        rows = table.get("rows", [])
        if not rows: continue

        first_row_cells = rows[0].get("cells", [])
        if len(first_row_cells) < 2: continue

        first_row_text = " ".join(str(c).lower() for c in first_row_cells)
        
        if any(kw in first_row_text for kw in ["назва банку", "надходження", "оборот", "тис. грн"]):
            # Пропускаємо 2 рядки (основний заголовок та його дубль)
            for row in rows[2:]: 
                cells = row.get("cells", [])
                if len(cells) >= 2:
                    turnover_data.append(cells[:2])
            break

    # === 4. Фільтрація числових значень ===
    for row in turnover_data:
        bank_name = str(row[0]).strip()
        value_str = str(row[1]).strip().replace(" ", "")

        numeric = re.sub(r"[^\d.,]", "", value_str)
        numeric = re.sub(r"\.+", ".", numeric)

        if numeric and any(char.isdigit() for char in numeric):
            numeric = numeric.replace(",", ".")
            clean_turnover.append([bank_name, numeric])

    # === 5. Розрахунки та застосування правил (Rules) ===
    if not clean_turnover:
        p_empty = doc.add_paragraph("Дані про обороти по рахункам відсутні")
        p_empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        financial_data["total_inflow"] = 0.0

        print("Diagnostics: No numeric turnover data")
    else:
        total_inflow = sum(float(val) for _, val in clean_turnover)
        financial_data["total_inflow"] = float(total_inflow)  

        # Розрахунок ratio: ділимо на revenue/1000, бо виручка зазвичай у тис. грн
        ratio = (total_inflow / (revenue / 1000) * 100) if revenue > 0 else 0

    # === ПЕРЕВІРКА УМОВИ < 50% (ІНТЕГРАЦІЯ В РЕЗУЛЬТАТ) ===
        if ratio < 50:
            cash_comment = "⚠ Переважають готівкові розрахунки (низька частка офіційних оборотів)"
            
            # Ініціалізуємо список, якщо його ще немає
            if "Мінуси" not in result:
                result["Мінуси"] = []
            
            # Додаємо коментар, якщо його там ще немає
            if cash_comment not in result["Мінуси"]:
                result["Мінуси"].append(cash_comment)
                print(f"📊 Логічне правило додано в Мінуси: {cash_comment} ({ratio:.2f}%)")

        # === Побудова таблиці в Word ===
        turnover_table = doc.add_table(rows=len(clean_turnover) + 2, cols=3)
        set_table_borders(turnover_table)
                
        headers = ["Назва Банку", "Надходження на рахунок за останні 12 місяців (тис. грн.)", "Відношення оборотів до виручки клієнта"]
        for j, hdr in enumerate(headers):
            cell = turnover_table.rows[0].cells[j]
            cell.text = hdr
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(hdr)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Заповнення даними
        for i, (bank, val) in enumerate(clean_turnover, start=1):
            row_cells = turnover_table.rows[i].cells
            row_cells[0].text = bank
            row_cells[1].text = f"{float(val):,.2f}".replace(",", " ")
            row_cells[2].text = ""

        # Рядок "Всього"
        total_row = turnover_table.rows[len(clean_turnover) + 1].cells
        total_row[0].text = "Всього"
        total_row[1].text = f"{total_inflow:,.2f}".replace(",", " ")
        total_row[2].text = f"{ratio:.2f}%"
        
        for cell in total_row:
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        print(f"Diagnostics: Added turnover table. Revenue: {revenue}, Ratio: {ratio:.2f}%")
        doc.add_paragraph("")
    # ----------------------------------------------------------------------
    # ДІАГНОСТИКА: показуємо всі ключі
    # ----------------------------------------------------------------------
    print("\n🔍 ДІАГНОСТИКА json_data:")
    for key in json_data.keys():
        print(f"  → '{key}': {type(json_data[key])}")

    print("\n🔍 ДІАГНОСТИКА financial_data:")
    for key in financial_data.keys():
        print(f"  → '{key}': {financial_data[key]}")
    
    # ----------------------------------------------------------------------
    # Коротка інформація про бізнес позичальника
    # ----------------------------------------------------------------------
    doc.add_paragraph("")
    p = doc.add_paragraph("Коротка інформація про бізнес позичальника:")
    p.runs[0].bold = True
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)

    # === ПОШУК У ТАБЛИЦЯХ ===
    tables = json_data.get("tables", [])

    def extract_from_tables(section_title_keywords, default="Дані відсутні."):
        for table in tables:
            rows = table.get("rows", [])
            for i, row in enumerate(rows):
                # Отримуємо текст усіх комірок у рядку
                cells = [str(c).strip() for c in row.get("cells", []) if c]
                full_row_text = " ".join(cells).strip()

                for kw in section_title_keywords:
                    # Шукаємо ключове слово (без врахування регістру)
                    if kw.lower() in full_row_text.lower():
                        
                        # Визначаємо, чи є в цьому рядку ще щось, крім заголовка
                        # Прибираємо саме ключове слово з тексту для перевірки
                        content_after_kw = re.sub(re.escape(kw), "", full_row_text, flags=re.IGNORECASE).strip()
                        content_after_kw = re.sub(r"^[.,\s:–-]+", "", content_after_kw) # чистимо знаки

                        # Якщо після ключового слова майже нічого немає (коротше 3 символів)
                        # значить відповідь у наступному рядку таблиці
                        if len(content_after_kw) < 3:
                            if i + 1 < len(rows):
                                next_row_cells = [str(c).strip() for c in rows[i+1].get("cells", []) if c]
                                res = " ".join(next_row_cells).strip()
                                if res: return res
                            continue # якщо наступного рядка немає, шукаємо далі
                        
                        # Якщо ж текст довгий — повертаємо його, вичистивши заголовок
                        return content_after_kw

        return default

    # === 1. УНІВЕРСАЛЬНА ФУНКЦІЯ ОТРИМАННЯ ДАНИХ ===
    # === 1. ПІДГОТОВКА ТЕКСТОВОГО КОРПУСУ ДЛЯ ПЛАНУ Б ===
    all_table_text = ""
    for table in tables:
        for row in table.get("rows", []):
            all_table_text += " ".join([str(c) for c in row.get("cells", []) if c]) + " "
    
    # Єдиний масив тексту для пошуку "каші"
    full_text_corpus = " ".join(all_paragraphs) + " " + all_table_text
    full_text_corpus = " ".join(full_text_corpus.split()) # Нормалізація пробілів

    # === 2. КОНФІГУРАЦІЯ ЯКОРІВ ДЛЯ ПЛАНУ Б ===
    fallback_config = {
        "Бізнес": {
            "patterns": [
                r"1\.?\s*Опис\s+структури\s+основних\s+видів\s+діяльності",
                "Має магазин", "Продає", "КВЕД", 
                r"власн[ийаое]\s+магазин", r"орендован[ийаое]\s+магазин", r"оренда\s+магазину"
            ],
            "stop": "2. Опис власної / орендованої МТБ"
        },
        "МТБ": {
            "patterns": [
                r"2\.?\s*Опис\s+власної\s*/\s*орендованої\s+МТБ",
                "Має власний будинок", "Авто", "Площа", "кв.м",
                r"власн[ийаое]\s+нерухомість", r"орендован[ийаое]\s+приміщення",
                r"власн[ийаое]\s+авто", r"орендован[ийаое]\s+авто",
                r"власн[ийаое]\s+ТЗ", r"орендован[ийаое]\s+ТЗ",
                r"оренда\s+нерухомості", r"орендує\s+приміщення"
            ],
            "stop": "3. Щодо поточного стану бізнесу"
            },
        # Розділ Воєнного стану по шматочках
        "Мобілізація": {
            "patterns": [r"Відсутні\s+проблеми\s+в\s+бізнесі", r"Повістку\s+не\s+отримував", r"Вік\s+позичальника.*?мобілізації", r"відсутні\s+проблеми"],
            "stop": "наявність/відсутність"
        },
        "Втрати_та_Бізнес": {
            "patterns": [
            r"Відсутні\s+проблеми\s+в\s+бізнесі",     
            r"Відсутні\s+проблеми", 
            r"Проблеми\s+відсутні",
            r"наявність/відсутність\s+суттєвих\s+втрат",
            r"можливість\s+подальшого\s+ведення\s+бізнесу"
            ],
            "stop": "поточний стан розрахунків"
        },
        "Розрахунки_та_Запаси": {
            "patterns": [r"Відсутні\s+проблеми\s+в\s+бізнесі", r"поточний\s+стан\s+розрахунків", r"рівень\s+запасів", r"можливість\s+вивозу\s+запасів", r"відсутні\s+проблеми"],
            "stop": "інша інформація"
        },
        "Обслуговування_Кредитів": {
            "patterns": [r"Відсутні\s+проблеми\s+в\s+бізнесі", r"Стан\s+обслуговування\s+кредитів", r"інша\s+інформація", r"відсутні\s+проблеми"],
            "stop": "Працює під час відсутності"
        },
        "Енергонезалежність": {
        "patterns": [r"Відсутні\s+проблеми\s+в\s+бізнесі", 
            r"Працює\s+під\s+час\s+відсутності.*?електроенергії", 
            r"наявний\s+потужний\s+генератор", r"відсутні\s+проблеми",
            r"придбати\s+обладнання.*?безперебійно"
        ],
        "stop": "IV."
        },
        "Динаміка_виручки": {
            "patterns": [
                r"1\.?\s*Причини\s+зменшення/збільшення\s+виручки",
                r"Причини\s+зменшення/збільшення\s+виручки"
            ],
            "stop": "2. Напрями виведення"
        },
        "Виведення_капіталу (за наявності)": {
            "patterns": [
                r"2\.?\s*Напрями\s+виведення\s+власного\s+капіталу",
                r"Напрями\s+виведення\s+власного\s+капіталу",
                "не відбувалося", "не виводив", "не відбувалось"
            ],
            "stop": "3. Напрями використання"
        },
        "Використання_прибутку": {
            "patterns": [
                r"3\.?\s*Напрями\s+використання\s+чистого\s+прибутку",
                r"Напрями\s+використання\s+чистого\s+прибутку",
                "накопичується", "використовується", "накопичуються", 
                "для розширення товарного асортименту", "особисті цілі"
            ],
            "stop": "4. Основні покупці"
        },
    # Контрагенти теж окремо
        "Покупці_Детально": {
            "patterns": [r"Покупці\s*[-–]\s*фіз\.особи", r"Покупці\s*[-–]"],
            "stop": "Постачальники"
        },
        "Постачальники_Детально": {
            "patterns": [r"Постачальники\s+Виробничо[- ]торг", r"Постачальники\s*[-–]"],
            "stop": "V."
        }
    }
    def get_data(key, table_keywords):
        res = "Дані відсутні."
        
        # А) Таблиці
        table_res = extract_from_tables(table_keywords, default=None)
        if table_res: res = table_res
        
        # Б) ParaMap
        if (res == "Дані відсутні." or not res) and key in para_map:
            res = para_map.get(key, "Дані відсутні.")
        
        # В) ПЛАН Б (Покращений)
        if (res == "Дані відсутні." or len(str(res)) < 5) and key in fallback_config:
            conf = fallback_config[key]
            
            for pattern in conf["patterns"]:
                # Гнучкий пошук: якщо є stop_marker — до нього, якщо ні — беремо шматок тексту
                if conf.get('stop'):
                    stop_part = re.escape(conf['stop'])
                    regex = rf"{pattern}[\s:–-]*(.*?)(?={stop_part})"
                else:
                    regex = rf"{pattern}[\s:–-]*(.{1,1500})" # Беремо наступні 1500 символів

                match = re.search(regex, full_text_corpus, re.IGNORECASE | re.DOTALL)
                
            if match:
                chunk = match.group(1).strip()
                
                # Покращена чистка: видаляємо лише технічне запитання в дужках
                # але залишаємо власне відповідь
                chunk = re.sub(r"\(У випадку наявності проблем зазначити\)", "", chunk, flags=re.IGNORECASE)
                
                # Якщо шматок починається з "в бізнесі", значить ми випадково 
                # відрізали "Відсутні проблеми". Додаємо їх назад або не видаляємо паттерн.
                if chunk.lower().startswith("в бізнесі"):
                        chunk = "Відсутні проблеми " + chunk
                
                if len(chunk) > 3:
                    return chunk.strip()
            if key in ["Воєнний_стан", "Контрагенти"]:
                print(f"--- DEBUG {key} ---")
                print(f"Результат перед поверненням: {res[:100]}...")
        return res
    # === 2. ПІДГОТОВКА ПАРАГРАФІВ (виправлена логіка) ===
    para_map = {}
    if all_paragraphs:
        for i in range(len(all_paragraphs) - 1):
            curr = " ".join(all_paragraphs[i].split()).strip() # Чистимо від \xa0 та пробілів
            nxt = all_paragraphs[i+1].strip()
            
            # Використовуємо 'in' для гнучкості
            if "Ціль кредитування" in curr:
                para_map["Ціль"] = nxt
            elif "Досвід роботи" in curr:
                para_map["Власники"] = nxt
            elif "Опис структури основних видів" in curr:
                para_map["Бізнес"] = nxt
            elif "Опис власної / орендованої МТБ" in curr:
                para_map["МТБ"] = nxt
            elif "Щодо поточного стану бізнесу" in curr:
                para_map["Воєнний_стан"] = nxt
            # ... інші elif залишаються такими ж, але використовуйте 'in' замість 'startswith'

    # === 3. ФОРМУВАННЯ ДОКУМЕНТА (Структуровано) ===

    # --- I. Опис угоди ---
    doc.add_heading("I. Опис угоди", level=1)
    target_value = get_data("Ціль", ["1. Ціль кредитування"])

    p = doc.add_paragraph()
    p.add_run("Ціль кредитування:").bold = True
    # Видаляємо можливе дублювання фрази "постачальник, власний внесок" з результату
    clean_target = re.sub(r"постачальник, власний внесок, тощо" or r"Втрат від бойових ", "", target_value, flags=re.IGNORECASE).strip()
    p.add_run(clean_target.lstrip(', :'))

    # --- II. Власники ---
    doc.add_heading("II. Власники та структура групи", level=2)
    ownership_text = get_data("Власники", ["1. Досвід роботи в даній сфері"])
    # Прибираємо дублювання довгої фрази заголовка з тексту
    clean_ownership = re.sub(r"види діяльності та взаємодія між членами.*компаній", "", ownership_text, flags=re.IGNORECASE).strip()
    doc.add_paragraph(clean_ownership.lstrip(', :'))


    # === III. Опис бізнесу ===
    doc.add_heading("III. Опис бізнесу", level=1) # Або level=3, залежно від вашої ієрархії

    # --- 3.1 Структура діяльності ---
    p = doc.add_paragraph("Опис структури основних видів діяльності (вказувати КВЕД не потрібно).")
    try:
        p.style = "List Number"
    except KeyError:
        p.style = "Normal"
    business_text = get_data("Бізнес", ["1. Опис структури основних видів діяльності (вказувати КВЕД не потрібно).", "Структура діяльності"])
    # Вичищаємо можливе дублювання інструкції з таблиці
    business_text = re.sub(r"\(вказувати КВЕД не потрібно\)\.?\s*", "", business_text, flags=re.IGNORECASE).strip()
    doc.add_paragraph(business_text.lstrip('., :–'))

    # --- 3.2 МТБ ---
    p = doc.add_paragraph("Опис власної / орендованої МТБ (нерухомість, транспорт, техніка та обладнання).")
    try:
        p.style = "List Number"
    except KeyError:
        p.style = "Normal"
    # Отримуємо дані
    mtb_raw = get_data("МТБ", ["Опис власної / орендованої МТБ", "МТБ"])
    # Вичищаємо довгий хвіст заголовка, який часто копіюють у комірку
    mtb_clean = re.sub(r"\(нерухомість, транспорт, техніка та обладнання\)?.*", "", mtb_raw, flags=re.IGNORECASE).strip()
    doc.add_paragraph(mtb_clean.lstrip('., :–'))

    # === III. 3.3 ВОЄННИЙ СТАН (ЗБІРНИЙ) ===
    doc.add_paragraph("Щодо поточного стану бізнесу в умовах воєнного стану та впливу мобілізації.")
    try:
        p.style = "List Number"
    except KeyError:
        p.style = "Normal"
    war_keys = [
        ("Мобілізація", ["Повістка", "Мобілізація"]),
        ("Втрати_та_Бізнес", ["Втрати", "Ведення бізнесу", "Пошкодження майна", "Пошкодження"]),
        ("Розрахунки_та_Запаси", ["Розрахунки", "Запаси"]),
        ("Обслуговування_Кредитів", ["Обслуговування", "Інша інформація"]),
        ("Енергонезалежність", ["Генератор", "Електроенергія"])
    ]
    
    combined_war_text = []
    seen_texts = set()  # Множина для відстеження вже доданих фрагментів

    for f_key, t_keywords in war_keys:
        chunk = get_data(f_key, t_keywords)
        
        if chunk and chunk != "Дані відсутні." and len(chunk) > 3:
            # Очищуємо від залишків паттернів
            clean_chunk = re.sub(r"^.*?–\s*", "", chunk).strip()
            
            # ПЕРЕВІРКА НА ДУБЛІКАТИ (використовуємо хеш або частину тексту)
            # Ми беремо перші 50 символів для порівняння
            text_fingerprint = clean_chunk[:50].lower()
            
            if text_fingerprint not in seen_texts:
                combined_war_text.append(clean_chunk)
                seen_texts.add(text_fingerprint)
                print(f"✅ Додано унікальний блок ({f_key})")
            else:
                print(f"⚠️ Пропущено дублікат для ключа: {f_key}")

    if combined_war_text:
        # Робимо текст зв'язним
        final_text = " ".join(combined_war_text)
        doc.add_paragraph(final_text)
    else:
        doc.add_paragraph("Дані щодо впливу воєнного стану не знайдені або бізнес працює стабільно.")

        
    # --- IV. Фінансовий аналіз ---
    doc.add_heading("IV. Фінансовий аналіз", level=2)

    sections_fin = [
        ("Причини зміни виручки", "Динаміка_виручки", ["Причини зменшення/збільшення виручки"]),
        ("Виведення капіталу (за наявності)", "Виведення_капіталу", ["Напрями виведення власного капіталу"]),
        ("Використання прибутку", "Використання_прибутку", ["Напрями використання чистого прибутку"]),
        ("Контрагенти", "Контрагенти", ["Основні покупці/постачальники"])
    ]

    for label, map_key, table_kw in sections_fin:
        p = doc.add_paragraph(label)
        try:
            p.style = "List Number"
        except KeyError:
            p.style = "Normal"
        doc.add_paragraph(get_data(map_key, table_kw))

    # === Діючі кредити в інших банках (Три колонки за задумкою) ===
    doc.add_heading("Платежі по кредитах клієнта та членів ГПК в інших банках:", level=2)
    
    credit_table = doc.add_table(rows=1, cols=3)
    set_table_borders(credit_table)
    credit_table.autofit = True
    
    # Ваші заголовки
    headers = ["Позичальник / Код ГПК", "Сума кредитів, тис. грн", "Річні платежі, тис. грн"]
    
    for i, header in enumerate(headers):
        cell = credit_table.rows[0].cells[i]
        cell.text = header
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 1. Отримуємо дані про платежі (вони вже згруповані по людях/кодах)
    borrower_payments = credit_data_res.get("payments_by_borrower", {})
    
    borrower_limits = {}
    details = credit_data_res.get("calculated_credit_details", [])
    for item in details:
        # Очищуємо ключ: прибираємо дужки, пробіли та лапки
        raw_name = str(item.get("borrower", "Невідомо"))
        clean_name = raw_name.replace("(", "").replace(")", "").strip()
        
        limit_val = item.get("limit", 0)
        borrower_limits[clean_name] = borrower_limits.get(clean_name, 0) + limit_val

    # 2. Уніфікуємо платежі
    borrower_payments = {}
    raw_payments = credit_data_res.get("payments_by_borrower", {})
    for b_name, p_val in raw_payments.items():
        clean_name = str(b_name).replace("(", "").replace(")", "").strip()
        borrower_payments[clean_name] = borrower_payments.get(clean_name, 0) + p_val

    # 3. Тепер зводимо все в одну множину унікальних імен
    all_borrowers = sorted(set(list(borrower_limits.keys()) + list(borrower_payments.keys())))

    # 4. Заповнюємо таблицю (тепер рядків буде стільки, скільки унікальних людей)
    for borrower in all_borrowers:
        row_cells = credit_table.add_row().cells
        
        # Колонка 1: Очищений код/ім'я
        row_cells[0].text = borrower
        
        # Колонка 2: Сума лімітів
        l_sum = borrower_limits.get(borrower, 0) / 1000
        row_cells[1].text = f"{l_sum:,.1f}".replace(",", " ")
        
        # Колонка 3: Річні платежі
        p_sum = borrower_payments.get(borrower, 0) / 1000
        row_cells[2].text = f"{p_sum:,.1f}".replace(",", " ")
        
        # Форматування шрифтів у новому рядку
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cell.runs[0] if p_cell.runs else p_cell.add_run(p_cell.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    doc.add_paragraph("") # Відступ після таблиці
            
    # ----------------------------------------------------------------------
    # Андерайтинг критерії — виправлено 0% + коментарі
    # ----------------------------------------------------------------------
    doc.add_paragraph("")
    doc.add_heading("Андерайтинг критерії", level=1)
        
    # Дані
    data = [
        ("Показник ліквідності", 
            f"{financial_data.get('KL', 0):.2f}" if isinstance(financial_data.get('KL'), (int, float)) 
            else financial_data.get('KL', "0.00")),
        ("Достатність капіталу", f"{financial_data.get('ER', 0):.2%}"),
        ("Рентабельність діяльності", f"{financial_data.get('ROS', 0):.2%}"),
        ("Загальний ЛАО / виручка", f"{financial_data.get('LAO_to_revenue', 0):.2f}%"),
        ("Ліміт WCF / виручка", f"{financial_data.get('WCF', 0):.2f}%"),
        ("Незабезпечений ліміт в РБ", f"{financial_data.get('unsecured_limit', 0):,.0f}".replace(",", " ")),
        ("Незабезпечений ліміт в РБ/Виручка", f"{financial_data.get('unsecured_to_revenue', 0):.2f}%"),
        ("ЛАО / обороти 12 міс", f"{financial_data.get('LAO_to_turnover', 0):.2f}%"),
        ("ROA", f"{financial_data.get('ROA', 0):.2f}%"),
        ("DSCR", f"{financial_data.get('DSCR', 0):.2f}")
        ]

    # Таблиця
    table = doc.add_table(rows=1 + len(data), cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Коефіцієнт"
    hdr[1].text = "Результат"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (name, value) in enumerate(data, start=1):
        row = table.rows[i]
        row.cells[0].text = name
        row.cells[1].text = value
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Підсвітка
        if name == "Незабезпечений ліміт в РБ/Виручка" and financial_data.get('unsecured_to_revenue', 0) > 20:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

    # Коментар під таблицею
    if financial_data.get('unsecured_to_revenue', 0) > 20:
        p = doc.add_paragraph("Відхилення від КП: значення > 20%")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        p.runs[0].bold = True

        max_limit = financial_data.get('max_unsecured_limit', 0)
        p = doc.add_paragraph(f"Максимальний ліміт кредитування по незабезпеченому ліміті в РБ становить - {max_limit:,.0f}".replace(",", " "))
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        p.runs[0].bold = True

    # Пропозиція по продуктам
    if financial_data.get("DSCR", 0) < 1.2:
        doc.add_paragraph("")
        p = doc.add_paragraph("Пропозиція по продуктам при DSCR = 1.2:")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        for credit in financial_data.get("credits", []):
            product = credit.get("type", "Невідомий")
            original = credit.get("amount", 0)
            max_amount = credit.get("max_amount_dscr_1_2", 0)
            if max_amount < original:
                p = doc.add_paragraph(f"• {product}: {max_amount:,.0f} (було {original:,.0f})".replace(",", " "))
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            else:
                p = doc.add_paragraph(f"• {product}: {max_amount:,.0f}".replace(",", " "))
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].font.size = Pt(10)
    
    # Додаємо заголовок перед таблицею
    p = doc.add_paragraph()
    run = p.add_run("Інформація щодо реального стану Позичальника")
    run.bold = True
    run.font.size = Pt(11)
    
    # Створюємо таблицю 2xN (Показник та Значення)
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    metrics_to_show = [
        ("Виручка від реалізації", "revenue"),
        ("Собівартість", "costs"),
        ("EBITDA", "ebitda"),
        ("Власний капітал", "equity"),
        ("Всього баланс", "total_balance"),
        ("Всього обігові", "current_assets"),
        ("Основні засоби", "fixed_assets"),
        ("Короткострокові зобов'язання", "short_term_liabilities"),
        ("Власний капітал", "equity"),
        ("URG", "scoring_result")
        ]
    
    for label, key in metrics_to_show:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        
        # Отримуємо значення, форматуємо його (розділення тисяч)
        value = financial_data.get(key, 0)
        if isinstance(value, (int, float)):
            row_cells[1].text = f"{value:,.0f}".replace(",", " ")
        else:
            row_cells[1].text = str(value)
            
        # Налаштування шрифту для комірок
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.size = Pt(10)

    # Додаємо відступ після таблиці
    doc.add_paragraph()

    # ----------------------------------------------------------------------
    # ТАБЛИЦЯ ПРІОРИТЕТНОСТІ ПРОДУКТІВ
    # ----------------------------------------------------------------------

    priority_data = [
        ("1", "Розвиток"),
        ("2", "Інвест"),
        ("3", "ВКЛ або Довіра"),
        ("4", "ОВД"),
    ]   

    # 1. Збираємо весь витягнутий текст в одну змінну
    full_extracted_text = " ".join(para_map.values())
    try:
        with open("rules.json", "r", encoding="utf-8") as f:
            text_rules = json.load(f)
    except Exception as e:
        print(f"⚠️ Файл rules_text.json не знайдено, використовую порожні правила: {e}")
        text_rules = {} # Якщо файлу немає, створюємо порожній список
    # 2. Запускаємо пошук за правилами з rules.json
    if text_rules:
        # Аналізуємо Плюси
        text_pluses = get_text_comments(full_extracted_text, text_rules.get("Плюси", []))
        # Аналізуємо Ризики (у вашому коді вони йдуть в "Мінуси")
        text_minuses = get_text_comments(full_extracted_text, text_rules.get("Ризики", []))

        text_pluses = list(set(text_pluses))
        text_minuses = list(set(text_minuses))
        
        # Додаємо результати до об'єкта result
        if "Плюси" not in result: result["Плюси"] = []
        if "Мінуси" not in result: result["Мінуси"] = []
        
        result["Плюси"].extend(text_pluses)
        result["Мінуси"].extend(text_minuses)
        
        # Прибираємо дублікати
        result["Плюси"] = list(dict.fromkeys(result.get("Плюси", []) + text_pluses))
        result["Мінуси"] = list(dict.fromkeys(result.get("Мінуси", []) + text_minuses))

    # Text analysis results
    doc.add_heading("Результат текстового аналізу", level=1)
    for key in ["Плюси", "Мінуси"]:
        doc.add_heading(key, level=2)
        
        if not result.get(key):
            p = doc.add_paragraph("(немає)")
            try:
                p.style = "List Bullet"
            except KeyError:
                p.style = "Normal"
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(12)
        else:
            for item in result[key]:
                prefix = "+ " if key == "Плюси" else "- "
                p = doc.add_paragraph(f"{prefix}{item.lstrip('+ ').lstrip('- ')}")
                try:
                    p.style = "List Bullet"
                except KeyError:
                    p.style = "Normal"
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)
 
    # === 1. Аналіз показників (наповнення списку висновків) ===
    fin_analysis_results = []
    credit_data_res = extract_credit_payments_from_json(json_data)
    # Отримуємо деталі кредитів з нашого результату
    all_details = credit_data_res.get("calculated_credit_details", [])
    
    # Фільтруємо саме картки
    credit_cards = [c for c in all_details if "кк" in str(c.get("type", "")).lower()]
    cc_count = len(credit_cards)
    
    if cc_count > 5:
        try:
            # Рахуємо загальну суму лімітів по КК
            cc_total_sum = sum(
                float(str(c.get("limit", 0)).replace(" ", "").replace(",", ".")) 
                for c in credit_cards
            ) / 1000
            
            # Формуємо коментар як МІНУС
            cc_comment = f"Кількість КК ({cc_count}) перевищує 5 шт. на загальну суму {cc_total_sum:.0f} тис. грн."
            fin_analysis_results.append(("-", cc_comment))
        except Exception as e:
            print(f"Помилка при аналізі КК для фінального списку: {e}")

    # Аналіз ліквідності (KL)
    val_kl = financial_data.get('KL')
    if isinstance(val_kl, (int, float)):
        if val_kl >= 1.0:
            fin_analysis_results.append(("+", f"Ліквідність: Добра ({val_kl:.2f})"))
        else:
            fin_analysis_results.append(("-", f"Ліквідність: Низька ({val_kl:.2f})"))
    elif val_kl == "абсолютна ліквідність":
        fin_analysis_results.append(("+", "Ліквідність: Абсолютна (відсутність короткострокових зобов'язань)"))

    # Аналіз достатності капіталу (ER)
    val_er = financial_data.get('ER', 0)
    if val_er >= 0.5:
        fin_analysis_results.append(("+", f"Достатність капіталу: Висока ({val_er:.1%})"))
    elif val_er >= 0.2:
        fin_analysis_results.append(("+", f"Достатність капіталу: Задовільна ({val_er:.1%})"))
    else:
        fin_analysis_results.append(("-", f"Достатність капіталу: Низька ({val_er:.1%})"))

    # Аналіз рентабельності (ROS)
    val_ros = financial_data.get('ROS', 0)
    if val_ros >= 0.1:
        fin_analysis_results.append(("+", f"Рентабельність діяльності: Прийнятна ({val_ros:.1%})"))
    elif val_ros > 0:
        fin_analysis_results.append(("-", f"Рентабельність діяльності: Низька ({val_ros:.1%})"))
    else:
        fin_analysis_results.append(("-", "Рентабельність діяльності: Від'ємна (збиток)"))

    # Беззаставний ліміт
    unsec_to_rev = financial_data.get("unsecured_to_revenue", 0)
    if unsec_to_rev > 20:
        fin_analysis_results.append(("-", f"Відхилення: незабезпечений ліміт перевищує 20% виручки ({unsec_to_rev:.1f}%)"))

    doc.add_heading("Результати фінансового аналізу", level=1)

    # 2. Створення таблиці для оцінок (Плюси/Мінуси)
    # Ми використовуємо 2 колонки: Статус (іконка) та Опис показника
    fin_table = doc.add_table(rows=0, cols=2)
    
    # Налаштування ширини колонок (опціонально)
    fin_table.columns[0].width = Inches(0.5)
    fin_table.columns[1].width = Inches(5.5)

    # 3. Заповнюємо таблицю даними з fin_analysis_results
    # fin_analysis_results — це список кортежів [("✓", "Текст..."), ("⚠", "Текст...")]
    for icon, description in fin_analysis_results:
        row_cells = fin_table.add_row().cells
        row_cells[0].text = icon
        row_cells[1].text = description
        
        # Можна додати форматування шрифту для іконок
        row_cells[0].paragraphs[0].alignment = 1 # Центрування

    # Main contractors
    # === 1. ДИНАМІЧНИЙ ПОШУК ТАБЛИЦІ ЗА КЛЮЧЕМ "TaxCode контрагента" ===
    contr_data = []
    
    for table in json_data.get('tables', []):
        rows = table.get('rows', [])
        if not rows: continue
        
        # Об'єднуємо текст перших двох рядків (шапки), видаляючи пробіли та переноси
        # Це допоможе знайти "TaxCodeконтрагента" навіть якщо вони розбиті
        header_area_text = ""
        for i in range(min(len(rows), 3)): # Перевіряємо перші 3 рядки про всяк випадок
            cells = rows[i].get('cells', [])
            header_area_text += "".join([str(c.get('text', '') if isinstance(c, dict) else c) for c in cells])
        
        clean_header = re.sub(r'\s+', '', header_area_text).lower()
        
        # Перевіряємо наявність ключа (без урахування пробілів)
        if "taxcodeконтрагента" in clean_header or "єдрпоуконтрагента" in clean_header:
            
            # Визначаємо, з якого рядка починаються дані. 
            # Зазвичай це наступний рядок після того, де знайшли слово TaxCode
            start_row_idx = 0
            for i, row in enumerate(rows[:3]):
                row_text = "".join([str(c.get('text', '') if isinstance(c, dict) else c) for c in row.get('cells', [])]).lower()
                if "taxcode" in row_text:
                    start_row_idx = i + 1
                    break

            # Збираємо дані
            for row in rows[start_row_idx:]:
                cells = row.get('cells', [])
                row_values = [str(c.get('text', '').strip() if isinstance(c, dict) else str(c).strip()) for c in cells]
                
                # Умови зупинки та фільтрації
                row_combined = " ".join(row_values).lower()
                if not any(row_values) or "разом" in row_combined or "всього" in row_combined:
                    continue
                
                # Додаємо перші 3 колонки (Код, Назва, %)
                if len(row_values) >= 2:
                    contr_data.append(row_values[:3])
            
            if contr_data: break

    # 2. ВИДАЛЯЄМО АБО КОРЕГУЄМО "ЗАПАСНИЙ ВАРІАНТ"
    # Якщо хочете бачити ПІБ, коли немає контрагентів — залишайте. 
    # Якщо хочете порожню таблицю — замініть на contr_data = [["-", "Дані відсутні", "-"]]
    if not contr_data:
        c_name = json_data.get('contractor_name')
        if c_name:
            contr_data = [[json_data.get('contractor_code', ''), c_name, "100"]]

    # 3. ТЕПЕР СТВОРЮЄМО ТАБЛИЦЮ У WORD
    doc.add_heading("Основні контрагенти", level=1)
    contr_table = doc.add_table(rows=1, cols=3)
    set_table_borders(contr_table)
    # Заголовки
    headers = ["TaxCode контрагента", "Найменування контрагента", "Частка контрагента %"]
    for i, h in enumerate(headers):
        cell = contr_table.rows[0].cells[i]
        cell.text = "" # Очистка
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    # Заповнення даними (БЕЗ ДУБЛЮВАННЯ)
    for data in contr_data:
        row_cells = contr_table.add_row().cells
        for i in range(3):
            val = str(data[i]) if i < len(data) else ""
            row_cells[i].text = "" # Очищуємо клітинку перед add_run
            run = row_cells[i].paragraphs[0].add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Блок пошуку та перепризначення таблиці параметрів ---
    target_table_id = "100"
    found_param_table = False

    for table in json_data.get("tables", []):
        # Отримуємо ID таблиці (перетворюємо на int для порівняння, якщо можливо)
        try:
            current_tid_str = table.get("table_id", "0")
            # Витягуємо тільки цифри, щоб порівняти "номер після 10"
            tid_numeric = int(re.search(r'\d+', current_tid_str).group()) if re.search(r'\d+', current_tid_str) else 0
        except (ValueError, AttributeError):
            tid_numeric = 0

        # Шукаємо тільки в таблицях з номером > 10
        if tid_numeric > 10:
            rows = table.get("rows", [])
            if not rows:
                continue
                
            # Об'єднуємо текст перших двох рядків для пошуку заголовків
            header_text = " ".join(" ".join(str(c).lower() for c in r.get("cells", [])) for r in rows[:2])
            
            # Перевіряємо наявність трьох ключових параметрів
            keywords = ["наименование параметра", "значение параметра", "сообщение"]
            if all(kw in header_text for kw in keywords):
                # Присвоюємо таблиці новий ID
                table["table_id"] = target_table_id
                found_param_table = True
                print(f"✅ Знайдено таблицю параметрів (оригінальний ID: {current_tid_str}). Присвоєно №100.")
                break # Виходимо, як тільки знайшли потрібну

    if not found_param_table:
        print("⚠️ Таблицю з параметрами (Наименование/Значение/Сообщение) не знайдено.")
# --------------------------------------------------------

    # 1. ПІДГОТОВКА (Виносимо словник на самий початок)
    translation = {
        "Наименование параметра": "Розрахункові параметри",
        "Значение параметра": "Значення параметру",
        "Сообщение": "Повідомлення"
    }
    calc_params_table = []
    found_by_id = False
    
    # Список ID, які ми категорично ігноруємо (2, 3, 4, 5)
    ignored_ids = ["Table_2", "Table_3", "Table_4", "Table_5", "Table_6", "Table_7", "Table_8", "Table_9", "Table_10", "Table_11"]

    try:
        tables = json_data.get('tables', [])
        
        # КРОК 1: Шукаємо суворо Table_13 (найвищий пріоритет)
        for table in tables:
            t_id = table.get("table_id")
            if table.get("table_id") == "100":
                rows = table.get('rows', [])
                calc_params_table = [[(c.get('text', '').strip() if isinstance(c, dict) else str(c).strip()) 
                                        for c in r.get('cells', [])] for r in rows if r.get('cells')]
                found_by_id = True
                print("🎯 Знайдено таблицю за пріоритетним ID: Table_13")
                break

        # КРОК 2: Якщо Table_13 не знайдено, шукаємо за назвою заголовка
        if not found_by_id:
            for table in tables:
                t_id = table.get("table_id")
                
                # Ігноруємо таблиці 2-5
                if t_id in ignored_ids:
                    continue
                
                rows = table.get('rows', [])
                if not rows: continue
                
                # Перевіряємо заголовок першої комірки
                first_cell = rows[0].get('cells', [])
                if first_cell:
                    c = first_cell[0]
                    text = (c.get('text', '') if isinstance(c, dict) else str(c)).lower()
                    
                    if "наименование" in text and "параметра" in text:
                        calc_params_table = [[(c.get('text', '').strip() if isinstance(c, dict) else str(c).strip()) 
                                                for c in r.get('cells', [])] for r in rows if r.get('cells')]
                        print(f"🔍 Table_13 відсутня. Знайдено за назвою заголовка в {t_id}")
                        break
                        
    except Exception as e:
        print(f"⚠️ Помилка при пошуку таблиці параметрів: {e}")
    if calc_params_table:
        doc.add_paragraph() 
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("Сірі зони")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 102, 204)

        # Створюємо таблицю один раз
        num_cols = len(calc_params_table[0]) if calc_params_table else 3
        table = doc.add_table(rows=0, cols=num_cols)
        set_table_borders(table)

        for row_data in calc_params_table:
            row_cells = table.add_row().cells
            full_row_text = " ".join(map(str, row_data)).lower()
            
            for i, cell_value in enumerate(row_data):
                if i < num_cols:
                    # Перекладаємо текст, якщо він є в словнику
                    val_str = str(cell_value).strip()
                    text_to_write = translation.get(val_str, val_str)
                    
                    # Очищуємо і записуємо
                    row_cells[i].text = "" 
                    p = row_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    run = p.add_run(text_to_write)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

                    # Логіка кольорів:
                    # Якщо це заголовок - робимо жирним
                    if "назва параметру" in text_to_write.lower() or "наименование" in val_str.lower():
                        run.bold = True
                    
                    # Якщо є негатив - червоним
                    if "негатив" in full_row_text or "не соответствует" in full_row_text :
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.bold = True
                    # 3. Виділення ЖОВТИМ маркером (Специфічні умови та Grey Zone)
                    # Створюємо список ключових слів для жовтого виділення
                    yellow_keywords = [
                        "grey zone", 
                        "енергоефективності", 
                        "більше ніж вказано в анкеті", 
                        "ручной расчет лимита",
                        "жовтих",
                        "ручний розрахунок ліміту"
                    ]

                    # Перевіряємо умови
                    if any(word.lower() in val_str.lower() or word.lower() in text_to_write.lower() for word in yellow_keywords):
                        # Отримуємо об'єкт поточного рядка. 
                        # Припустимо, ви зараз працюєте з row_cells (як у попередніх прикладах)
                        for cell in row_cells:
                            # 1. Робимо заливку всієї комірки (жовтий колір HEX: FFFF00)
                            set_cell_background(cell, "FFFF00") 
                            
                            # 2. Робимо текст жирним для всіх параграфів у цій комірці
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
    # Collateral
    doc.add_heading("Забезпечення по кредиту:", level=1)
    doc.add_paragraph(json_data.get('collateral', 'Забезпечення відсутнє'))
    
    # Other comments
    doc.add_heading("Інші коментарі:", level=1)
    doc.add_paragraph(json_data.get("comments", "Коментарі відсутні"))
    
    # === РОЗДІЛ: ПРОПОЗИЦІЯ (PROPOSAL) ===
    doc.add_heading("На підставі проведеного аналізу пропоную затвердити кредит на наступних умовах:", level=1)
    
    # Отримуємо значення DSCR для перевірки умови (припускаємо, що воно є в financial_data)
    current_dscr = financial_data.get('dscr', 1.2) 
    
    for index, request in enumerate(credit_requests, start=1):
        # 1. Визначаємо суму на основі DSCR
# --- 1. ЛОГІКА ВИЗНАЧЕННЯ СУМИ (DSCR ПЕРЕВІРКА) ---
        original_amount = float(request.get('amount', 0))
        max_limit_dscr = float(request.get('max_amount_dscr_1_2', 0))
        
        # Якщо DSCR в нормі (>= 1.2), даємо original. 
        # Якщо низький, даємо мінімум між запитом і розрахунковим лімітом
        if current_dscr >= 1.2:
            final_amount = original_amount
            dscr_note = "Сума затверджена в повному обсязі (DSCR >= 1.2)"
        else:
            final_amount = min(original_amount, max_limit_dscr)
            dscr_note = f"Суму обмежено до ліміту DSCR 1.2 (DSCR = {current_dscr:.2f}, було {original_amount/1000:,.0f} тис.)"

        # --- 2. СТВОРЕННЯ ТАБЛИЦІ ПРОПОЗИЦІЇ ---
        prop_table = doc.add_table(rows=2, cols=5)
        set_table_borders(prop_table)
        prop_table.autofit = False

        # Заголовки (стилізація)
        headers = ["Сума (тис. грн)", "Валюта", "Термін (міс.)", "Ставка", "Продукт"]
        for i, header in enumerate(headers):
            cell = prop_table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold, run.font.name, run.font.size = True, 'Times New Roman', Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(3.0)

        # Рядок з даними
        row_cells = prop_table.rows[1].cells
        
        # Формуємо значення (якщо сума 0 — виведе "0")
        proposal_values = [
            "{:,.0f}".format(final_amount / 1000).replace(',', ' '), 
            "UAH", 
            str(request.get('term', '36')), 
            "Згідно КУАП", 
            request.get('type', 'Невідомий')
        ]

        for i, val in enumerate(proposal_values):
            cell = row_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name, run.font.size = 'Times New Roman', Pt(10)
            
            # Якщо сума 0 і DSCR низький — підсвічуємо червоним
            if i == 0 and final_amount == 0 and current_dscr < 1.2:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # --- 3. ПРИМІТКА (тільки якщо було обмеження) ---
        if final_amount < original_amount:
            p_note = doc.add_paragraph(dscr_note)
            if p_note.runs:
                run_n = p_note.runs[0]
                run_n.font.size, run_n.italic = Pt(9), True
                run_n.font.color.rgb = RGBColor(200, 0, 0) # Темно-червоний
            print(f"⚠️ Кредит №{index} обмежено: {final_amount} (DSCR: {current_dscr})")

        doc.add_paragraph("") # Відступ
    
    # Signature
    p = doc.add_paragraph(f"Ризик менеджер: {final_risk_manager} ________________ ()")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save document
    results_dir = "results"
    os.makedirs(results_dir, exist_ok=True)
    filepath = os.path.join(results_dir, f"Висновок_андерайтера_{timestamp}.docx")
    doc.save(filepath)
    print(f"📝 Saved conclusion: {filepath}")

def process_document (uploaded_file, analyst_name):
    # 1. Ініціалізація порожнього словника 
    financial_data = {}
    output = None
    global fin_analysis_results
    fin_analysis_results = []
    try:
        # 2. Відкриваємо документ
        doc = Document(uploaded_file)
        print(f"🚀 Початок обробки файлу: {uploaded_file}")

        # 3. Викликаємо інструменти 
        raw_text, json_data = load_doc_text(uploaded_file)
        # --- ОЧИЩЕННЯ ШАБЛОНУ ---
        # Видаляємо старий зміст, щоб залишити тільки порожній файл зі стилями
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        for t in doc.tables:
            t._element.getparent().remove(t._element)
        # ------------------------
        if not raw_text.strip():
            print("⚠️ Document text is empty")
            return financial_data
    
        scoring_value = extract_scoring_result(raw_text)
        financial_signals = []
        credit_payments_info = []
        credit_history_payments = {"credit_history_payments": []}
        financial_data = {}
        result = {"Плюси": [], "Мінуси": []}
        credit_payments = {"total_annual_payments_other_banks": 0.0, "payments_by_borrower": {}}

        # 1. Отримуємо дані кредитної історії ОДИН РАЗ (це повертає словник з усіма сумами та списком)
        history_results = extract_credit_history_payments(json_data)
        
        # 2. Робота з текстом та метриками
        cleaned_text, original_text = raw_text, raw_text
        print(f"📝 Text cleaned, length: {len(cleaned_text)}")
        
        extracted_metrics = extract_financial_metrics_from_json(json_data, alias_map)
        financial_data.update(extracted_metrics)
        
        # 3. Нормалізація та витяг стандартних платежів
        financial_data = normalize_financial_data(financial_data)
        credit_payments = extract_credit_payments_from_json(json_data)
        financial_data.update(credit_payments)
        financial_data["scoring_result"] = scoring_value if scoring_value else "Н/Д"

        # 4. Оновлюємо фінансові дані результатами кредитної історії
        financial_data.update({
            "total_annual_payments_history": history_results["total_annual_payments_history"],
            "WCF_RBUA_total": history_results["WCF_RBUA_total"],
            "IF_RBUA_total": history_results["IF_RBUA_total"],
            "credit_history_payments": history_results["credit_history_payments"]
        })

        # 5. Готуємо змінні для передачі в save_results_to_docx
        # Щоб функція save_results_to_docx не "падала" на .get(), передаємо їй весь словник
        credit_history_payments = history_results 
        
        # Отримуємо параметри розрахунку
        calc_params_table = get_calculation_params_from_json(json_data)

        # 6. Створюємо таблицю в документі, передаючи ТІЛЬКИ список кредитів (history_results["credit_history_payments"])      
        default_rate = 0.20
        default_term = 12
        for credit in financial_data.get('credits', []):
            if not isinstance(credit, dict):
                print(f"⚠️ Invalid credit structure: {credit}")
                continue
            amount = credit.get("amount", 0.0)
            credit_type = credit.get("type", "Невідомий")
            rate = credit.get("rate", default_rate)
            term = credit.get("term", default_term)
            if amount > 0 and rate >= 0 and term > 0:
                annual_payment = amount * rate if any(keyword in credit_type.lower() for keyword in ["овердрафт", "вкл", "довіра"]) else (amount / (term / 12)) + (amount * rate)
                credit["annual_payment"] = annual_payment
                credit_payments_info.append({
                    "type": credit_type,
                    "amount": amount,
                    "rate": rate,
                    "term": term,
                    "annual_payment": annual_payment
                })
                print(f"✅ Added to credit_payments_info: {credit_payments_info[-1]}")
        
        has_long_term_credit = any(
            credit.get("term", 0) > 12 and not any(keyword in credit.get("type", "").lower() for keyword in ["овердрафт", "вкл", "довіра"])
            for credit in financial_data.get('credits', [])
        )
      
        try:
            # Пріоритетність продуктів (менше число = ВИЩИЙ пріоритет)
            priority_map = {
            "Розвиток": 1,
            "Інвест": 2,
            "ВКЛ": 3,
            "Довіра": 3,
            "ОВД": 4,
            "Овердрафт": 4
            }

            credits = financial_data.get("credits", [])
            # Сортуємо: ВИЩИЙ ПРІОРИТЕТ ПЕРШИЙ
            sorted_credits = sorted(
            credits,
            key=lambda c: priority_map.get(c.get("type", ""), 99)
            )

            print(f"Продукти для зменшення (пріоритет): {[c.get('type') for c in sorted_credits]}")

            current_dscr = financial_data.get("DSCR", 0)
            ebitda = financial_data.get("ebitda", 0)

            if current_dscr < 1.2 and ebitda > 0:
                other_debt_service = (
                    financial_data.get("total_annual_payments_other_banks", 0) +
                    financial_data.get("total_annual_payments_history", 0)
                )
                target_debt_service = ebitda / 1.2
                available_for_request = target_debt_service - other_debt_service
                print(f"Доступно до запиту (Available for Request): {available_for_request:,.2f} грн")
            if available_for_request > 0:
                remaining_capacity = available_for_request
                for credit in sorted_credits:
                    term = credit.get("term", 36)
                    rate = credit.get("rate", 0.17)
                    original_amount = credit.get("amount", 0)
                    product_name = credit.get("type", "Невідомий")

                    monthly_rate = rate / 12
                    if monthly_rate > 0:
                        coeff = monthly_rate / (1 - (1 + monthly_rate) ** (-term))
                    else:
                        coeff = 1 / term

                    max_amount = remaining_capacity / coeff if coeff > 0 else 0
                    credit["max_amount_dscr_1_2"] = max(0, min(original_amount, max_amount))
                    remaining_capacity -= credit["max_amount_dscr_1_2"] * coeff
                    print(f"Для '{product_name}': max = {credit['max_amount_dscr_1_2']:,.0f} (було {original_amount:,.0f})")

                    if remaining_capacity <= 0:
                        break
            else:
                for credit in sorted_credits:
                    credit["max_amount_dscr_1_2"] = 0
                else:
                    for credit in sorted_credits:
                        credits["max_amount_dscr_1_2"] = credit.get("amount", 0)

                financial_data["credits"] = sorted_credits

        except Exception as e:
            print(f"Помилка в розрахунку DSCR: {e}")

            import traceback
            print(f"❌ Помилка при обробці фінансових даних або кредитів: {e}")
            print(traceback.format_exc())
        # Можна додати значення за замовчуванням, щоб код не зламався далі
        if "scoring_result" not in financial_data:
            financial_data["scoring_result"] = "Помилка обробки"
        # === Аналіз кредитних карток (КК) по ГПК ===
       
        credit_data_res = extract_credit_payments_from_json(json_data)
        all_details = credit_data_res.get("calculated_credit_details", [])

        print(f"Всього знайдено записів про кредити в інших банках: {len(all_details)}")
        
        # Фільтруємо саме картки
        credit_cards = [c for c in all_details if "кк" in str(c.get("type", "")).lower()]
        cc_count = len(credit_cards)
        print(f"З них ідентифіковано як 'Кредитні картки': {cc_count}")
        
        if cc_count > 5:
            print(f"УВАГА: Кількість карток ({cc_count}) перевищує ліміт (5). Розраховую суму...")
            try:
                # Очищуємо рядок від пробілів та ком, щоб float не видав помилку
                cc_total_sum = sum(
                    float(str(c.get("limit", 0)).replace(" ", "").replace(",", ".")) 
                    for c in credit_cards
                ) / 1000
                
                cc_text = f"По ГПК {cc_count} кредитних карток на загальну суму {cc_total_sum:.0f} тис. грн."
                fin_analysis_results.append(("⚠", cc_text))
                
                print(f"РЕЗУЛЬТАТ: {cc_text}")
            except Exception as e:
                print(f"ПОМИЛКА при розрахунку суми карток: {e}")
        else:
            print(f"Кількість карток ({cc_count}) в межах норми (<=10), додатковий ризик не додається.")
        
        print("-----------------------------------\n")

        hist_result = extract_credit_history_payments(json_data)
        financial_data["WCF_RBUA"] = hist_result["WCF_RBUA_total"]
        financial_data["IF_RBUA"] = hist_result["IF_RBUA_total"]

        # --- Ручний розрахунок ---
        current_assets = financial_data.get("current_assets", 0)
        short_term_liabilities = financial_data.get("short_term_liabilities", 1e-9)
        equity = financial_data.get("equity", 0)
        total_balance = financial_data.get("total_balance", 1e-9)
        ebitda = financial_data.get("ebitda", 0)*1000
        revenue = financial_data.get("revenue", 1e-9)*1000
        total_requested_credit = financial_data.get("total_requested_credit", 0)
        turnover_12m = financial_data.get("turnover_12m", revenue)
        current_inflow = float(financial_data.get('total_inflow', 0))
        other_banks_payments = financial_data.get("total_annual_payments_other_banks", 0)*1000
        history_payments = financial_data.get("total_annual_payments_history", 0)*1000
        if_otherbanks = financial_data.get("IF_otherbanks_total", 0)*1000
        wcf_otherbanks = financial_data.get("WCF_otherbanks_total", 0)*1000
        RBUA_IF = financial_data.get("IF_RBUA", 0)*1000
        RBUA_WCF = financial_data.get("WCF_RBUA", 0)*1000        
        annual_payment = financial_data.get("total_annual_payment", 0)
        fin_analysis_results = []
        
        # КЛ
        try:
            # Розрахунок КЛ
            financial_data["KL"] = current_assets / short_term_liabilities
            print(f"Calculated KL: {financial_data['KL']:.2f}")
        except ZeroDivisionError:
            # Якщо short_term_liabilities == 0
            print("Абсолютна ліквідність")
            financial_data["KL"] = "абсолютна ліквідність"
        val_kl = financial_data.get("KL")
        if isinstance(val_kl, (int, float)):
            if val_kl > 1.0:
                fin_analysis_results.append(("✓", f"Коефіцієнт ліквідності: Добра ліквідність ({val_kl:.2f})"))
            elif val_kl > 0.5:
                fin_analysis_results.append(("⚠", f"Коефіцієнт ліквідності: Задовільна ліквідність ({val_kl:.2f})"))
            else:
                fin_analysis_results.append(("⚠", f"Коефіцієнт ліквідності: Низька ліквідність ({val_kl:.2f})"))
        elif val_kl == "абсолютна ліквідність":
            fin_analysis_results.append(("✓", "Коефіцієнт ліквідності: Абсолютна ліквідність"))    

        # ER
        financial_data["ER"] = equity / (total_balance + (total_requested_credit / 1000))
        print(f"Calculated ER: {financial_data['ER']:.2f}")
        # Перевірка ER (Достатність капіталу)
        val_er = financial_data.get("ER", 0)
        if val_er > 0.5:
            fin_analysis_results.append(("✓", f"Достатність капіталу: Висока ({val_er:.2f})"))
        elif val_er > 0.2:
            fin_analysis_results.append(("✓", f"Достатність капіталу: Задовільна ({val_er:.2f})"))
        else:
            fin_analysis_results.append(("⚠", f"Достатність капіталу: Низька ({val_er:.2f})"))

        # ROS
        financial_data["ROS"] = ebitda / revenue
        print(f"Calculated ROS: {financial_data['ROS']:.2f}")

        # Перевірка ROS (Рентабельність)
        val_ros = financial_data.get("ROS", 0)
        if val_ros > 0.3:
            fin_analysis_results.append(("✓", f"Операційна рентабельність: Висока ({val_ros:.2f})"))
        elif val_ros > 0.1:
            fin_analysis_results.append(("✓", f"Операційна рентабельність: Задовільна ({val_ros:.2f})"))
        elif val_ros > 0:
            fin_analysis_results.append(("⚠", f"Операційна рентабельність: Низька ({val_ros:.2f})"))
        else:
            fin_analysis_results.append(("⚠", "Операційна рентабельність: Неприбуткова діяльність"))

        # ЛАО / виручка
        total_debt = total_requested_credit + if_otherbanks + wcf_otherbanks + RBUA_WCF + RBUA_IF
        financial_data["LAO_to_revenue"] = (total_debt/ revenue)*100  if revenue > 0 else 0
        print(f"Calculated ЛАО/виручка: {financial_data['LAO_to_revenue']:.2f}")
        
        # WCF
        wcf_credit = RBUA_WCF/2 + wcf_otherbanks
        financial_data["WCF"] = (wcf_credit/ revenue)*100  if revenue > 0 else 0
        print(f"Calculated WCF: {financial_data['WCF']:.2f}")

        # Беззаставний ліміт
        unsec_limit=total_requested_credit + RBUA_WCF
        financial_data["unsecured_limit"] = unsec_limit / 1000
        print(f"Set unsecured_limit: {financial_data['unsecured_limit']:,.0f}")

        # ROA
        financial_data["ROA"] =  (ebitda / total_balance)/10 if total_balance > 0 else 0
        print(f": {financial_data['ROA']:.2f}")
    
        # ЛАО / обороти
        # --- ПРИМУСОВЕ ВИЗНАЧЕННЯ ОБОРОТІВ (якщо total_inflow ще 0) ---
        if financial_data.get("total_inflow", 0) == 0:
            temp_inflow = 0.0
            # Шукаємо таблицю оборотів в json_data
            for table in json_data.get("tables", []):
                rows = table.get("rows", [])
                if not rows: continue
                
                # Перевіряємо, чи це таблиця оборотів
                first_row_text = " ".join(str(c).lower() for c in rows[0].get("cells", []))
                if any(kw in first_row_text for kw in ["назва банку", "надходження", "оборот"]):
                    # Сумуємо дані з другого стовпця (починаючи з 3-го рядка)
                    for row in rows[2:]:
                        cells = row.get("cells", [])
                        if len(cells) >= 2:
                            val_str = str(cells[1]).strip().replace(" ", "").replace(",", ".")
                            # Видаляємо все крім цифр та крапки
                            val_str = re.sub(r"[^\d.]", "", val_str)
                            if val_str:
                                try:
                                    temp_inflow += float(val_str)
                                except: continue
                    break # Знайшли потрібну таблицю, виходимо з циклу
            
            financial_data["total_inflow"] = temp_inflow
            print(f"📊 Примусово знайдено обороти: {temp_inflow}")

        # --- ТЕПЕР РОЗРАХУНОК ЛАО / ОБОРOTИ ---
        current_inflow = financial_data.get("total_inflow", 0)
        
        if current_inflow > 0:
            # Розрахунок у %
            financial_data["LAO_to_turnover"] = ((total_debt / 10000) / current_inflow)
            print(f"✅ Результат ЛАО/Обороти: {financial_data['LAO_to_turnover']:.2f}%")
        else:
            financial_data["LAO_to_turnover"] = 0
            print("⚠️ Не вдалося розрахувати ЛАО/Обороти: Обороти не знайдено")
        
        # DSCR
        total_debt_service = annual_payment + other_banks_payments + history_payments
        financial_data["DSCR"] = ebitda / total_debt_service if total_debt_service > 0 else 0
        print(f"Calculated DSCR: {financial_data['DSCR']:.2f}")

        # Беззаставний ліміт / Виручка
        revenue = financial_data.get("revenue", 1e-9)
        unsecured_limit = financial_data.get("unsecured_limit", 0)
        unsecured_to_revenue = (unsecured_limit / revenue)*100 if revenue > 0 else 0
        financial_data["unsecured_to_revenue"] = unsecured_to_revenue
        max_unsecured_limit = revenue * 0.2
        financial_data["max_unsecured_limit"] = max_unsecured_limit
        print(f"Незабезпечений ліміт / Виручка: {unsecured_to_revenue:.2f}%")
        if unsecured_to_revenue > 20:
            print("Відхилення від КП: значення > 20%")
            print(f"Максимальний ліміт кредитування по незабезпеченому ліміті в РБ становить - {max_unsecured_limit:,.0f}")
        
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

        # Формуємо повний шлях до файлу з фінансовими правилами
        finance_rules_path = os.path.join(BASE_DIR, "rules_finance.json")

        try:
            with open(finance_rules_path, "r", encoding="utf-8") as f:
                finance_rules = json.load(f)
            print(f"✅ Фінансові правила успішно завантажено: {len(finance_rules)} записів")
        except FileNotFoundError:
            print(f"❌ КРИТИЧНО: Файл не знайдено за шляхом: {finance_rules_path}")
            finance_rules = []
        except json.JSONDecodeError:
            print(f"❌ ПОМИЛКА: Файл rules_finance.json має некоректний формат (помилка в синтаксисі JSON)")
            finance_rules = []
        except Exception as e:
            print(f"❌ Непередбачена помилка при читанні правил: {e}")
            finance_rules = []
        # ----------------------------------------------------------------------
        # МАКСИМАЛЬНА СУМА КРЕДИТУ ПРИ DSCR = 1.2
        # ----------------------------------------------------------------------
        if financial_data.get("DSCR", 0) < 1.2 and financial_data.get("ebitda", 0) > 0:
            # Платежі без запиту
            other_debt_service = (
                financial_data.get("total_annual_payments_other_banks", 0) +
                financial_data.get("total_annual_payments_history", 0)
            )
            
            # Цільовий DSCR = 1.2
            target_ebitda_coverage = financial_data["ebitda"] * 1.2
            
            # Скільки можна додати на платіжі по запиту
            available_for_request = target_ebitda_coverage - other_debt_service
            
            if available_for_request > 0:
                # Припустимо, що термін і ставка — з першого кредиту
                first_credit = next((c for c in financial_data.get("credits", []) if c.get("term", 0) > 12), None)
                if first_credit:
                    term_months = first_credit.get("term", 36)
                    rate = first_credit.get("rate", 0.15)  # 15%
                    # Щомісячний платіж = (сума * (ставка/12)) / (1 - (1 + ставка/12)^(-термін))
                    monthly_rate = rate / 12
                    max_credit = available_for_request * (1 - (1 + monthly_rate) ** (-term_months)) / monthly_rate
                    financial_data["max_credit_dscr_1_2"] = max_credit
                    print(f"Calculated max_credit_dscr_1_2: {max_credit:,.0f}")
                else:
                    financial_data["max_credit_dscr_1_2"] = 0
            else:
                financial_data["max_credit_dscr_1_2"] = 0
        else:
            financial_data["max_credit_dscr_1_2"] = financial_data.get("total_requested_credit", 0)
            print(f"DSCR >= 1.2 — max_credit_dscr_1_2 = запит: {financial_data['max_credit_dscr_1_2']:,.0f}")
        
        try:
            with open("rules.json", "r", encoding="utf-8") as f:
                text_rules = json.load(f)
        except Exception as e:
            print(f"⚠️ Файл rules_text.json не знайдено, використовую порожні правила: {e}")
            text_rules = {} # Якщо файлу немає, створюємо порожній список
        for signal in financial_signals:
            name = signal.get("name", "")
            if name in alias_map:
                value = float(signal.get("value", 0))
                financial_data[alias_map[name]] = value
                print(f"Calculated {name}: {value:.2f}")
    
        result = evaluate_company(raw_text, text_rules, financial_signals)
        descriptions = {
        "Ціль кредитування": json_data.get("deal_description", ""),
        "Власники": json_data.get("ownership_info", ""),
        "Опис бізнесу": json_data.get("business_description", ""),
        "МТБ": json_data.get("mtb_info", "")
        }
        all_paragraphs = json_data.get("paragraphs", [])

        # --- КЛЮЧОВИЙ МОМЕНТ ЗБЕРЕЖЕННЯ ---
        save_results_to_docx(
            financial_signals, analyst_name, credit_data_res, result, datetime.now().strftime("%Y-%m-%d_%H-%M-%S"),
            credit_payments_info, credit_history_payments=history_results, json_data=json_data, financial_data=financial_data, 
            credit_payments=credit_payments, calc_params_table=calc_params_table, descriptions=descriptions, all_paragraphs=all_paragraphs, doc=doc
        )
        output = io.BytesIO()
        doc.save(output)
        output.seek(0) # Повертаємо "курсор" у початок файлу, щоб його можна було прочитати
        print(f"📊 ВИХІДНИЙ ФАЙЛ СФОРМОВАНО")

    except Exception as e:
        import traceback
        print("\n" + "!"*50)
        print(f"❌ КРИТИЧНА ПОМИЛКА ОБРОБКИ: {e}")
        traceback.print_exc() 
        print("!"*50 + "\n")
        
        # Якщо сталася помилка, додаємо її опис у словник
        if not financial_data:
            financial_data = {"error": str(e)}
        else:
            financial_data["processing_error"] = str(e)
            
    # Цей return стоїть на рівні з try/except, завершуючи функцію
    import gc
    import time
    gc.collect()  # Звільняє пам'ять
    start_cleanup = time.time()
    n_cleaned = gc.collect() 
    end_cleanup = time.time()
    print(f"🧹 Очищення пам'яті: видалено {n_cleaned} об'єктів за {end_cleanup - start_cleanup:.4f} сек.")
    return output, financial_data

if __name__ == "__main__":
    pass
